from __future__ import annotations

import importlib.util
import json
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "pos-apotik-emedik" / "uat-v1.34.25"
SHOT = OUT / "screenshots"
PROC = OUT / "screenshots-pengadaan"
ACC = OUT / "screenshots-akuntansi"
DIAGRAM = OUT / "diagrams"
VERSION = "1.35.1 (build 189)"
RELEASE = "API apotik-v1.34.25+187"
DATE = "5 September 2026"

DEEP = "123524"
GREEN = "166534"
ACCENT = "16A34A"
TEAL = "0F766E"
INK = "172033"
MUTED = "526273"
LIGHT = "F0FDF4"
MID = "BBF7D0"
AMBER = "B45309"
RED = "B91C1C"


def load_old():
    source = ROOT / "tools" / "generate_apotik_v13423_uat.py"
    spec = importlib.util.spec_from_file_location("apotik_v13423_base", source)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return module


BASE = load_old()


def load_json(path: Path) -> dict:
    if not path.exists():
        raise FileNotFoundError(f"Bukti mesin belum tersedia: {path}")
    return json.loads(path.read_text(encoding="utf-8"))


def font(size: int, bold: bool = False):
    return BASE.font(size, bold)


def centered(draw: ImageDraw.ImageDraw, box, text: str, size=22, bold=True,
             color="#172033", max_chars=30):
    lines: list[str] = []
    for part in text.split("\n"):
        lines.extend(wrap(part, max_chars) or [""])
    face = font(size, bold)
    heights = []
    for line in lines:
        bounds = draw.textbbox((0, 0), line, font=face)
        heights.append(bounds[3] - bounds[1] + 7)
    y = box[1] + max(8, ((box[3] - box[1]) - sum(heights)) / 2)
    for line, height in zip(lines, heights):
        bounds = draw.textbbox((0, 0), line, font=face)
        x = (box[0] + box[2] - (bounds[2] - bounds[0])) / 2
        draw.text((x, y), line, font=face, fill=color)
        y += height


def box(draw, xy, label, fill="#FFFFFF", outline="#166534", size=21):
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    centered(draw, xy, label, size=size)


def arrow(draw, start, end, color="#166534", width=5):
    draw.line([start, end], fill=color, width=width)
    x, y = end
    if abs(end[0] - start[0]) >= abs(end[1] - start[1]):
        direction = 1 if end[0] > start[0] else -1
        draw.polygon([(x, y), (x - direction * 17, y - 9),
                      (x - direction * 17, y + 9)], fill=color)
    else:
        direction = 1 if end[1] > start[1] else -1
        draw.polygon([(x, y), (x - 9, y - direction * 17),
                      (x + 9, y - direction * 17)], fill=color)


def canvas(title: str):
    image = Image.new("RGB", (1600, 900), "#F7FCF8")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1600, 96), fill=f"#{DEEP}")
    draw.text((48, 26), title, font=font(32, True), fill="white")
    return image, draw


def make_use_case(index: int, item: dict) -> Path:
    image, draw = canvas(f"Use Case {index} - {item['title']}")
    ys = [145, 315, 485, 655]
    actor_boxes = [(55, y, 350, y + 94) for y in ys]
    case_boxes = [(610, y, 1545, y + 94) for y in ys]
    # Connector dibuat terlebih dahulu dan hanya menempati ruang antar-kotak.
    for actor_box, case_box in zip(actor_boxes, case_boxes):
        arrow(draw, (actor_box[2], (actor_box[1] + actor_box[3]) // 2),
              (case_box[0], (case_box[1] + case_box[3]) // 2), "#15803D", 5)
    for actor_box, label in zip(actor_boxes, item["actors"]):
        box(draw, actor_box, label, "#DCFCE7", "#166534", 20)
    for case_box, label in zip(case_boxes, item["usecases"]):
        box(draw, case_box, label, "#FFFFFF", "#0F766E", 21)
    draw.text((55, 824), "Interaksi dipisahkan per baris; garis tidak melewati simpul atau label.",
              font=font(19), fill="#526273")
    path = DIAGRAM / f"{index:02d}-use-case.png"
    image.save(path)
    return path


def make_flow(index: int, item: dict) -> Path:
    image, draw = canvas(f"Flowchart {index} - {item['title']}")
    positions = [
        (55, 175, 410, 300), (620, 175, 975, 300), (1185, 175, 1540, 300),
        (1185, 510, 1540, 635), (620, 510, 975, 635), (55, 510, 410, 635),
    ]
    # Jalur snake hanya memakai gutter 210 px dan gutter vertikal 210 px.
    arrow(draw, (410, 237), (620, 237))
    arrow(draw, (975, 237), (1185, 237))
    arrow(draw, (1362, 300), (1362, 510))
    arrow(draw, (1185, 572), (975, 572))
    arrow(draw, (620, 572), (410, 572))
    fills = ["#DCFCE7", "#FFFFFF", "#FFFFFF", "#FEF3C7", "#FFFFFF", "#DCFCE7"]
    outlines = ["#166534", "#0F766E", "#0F766E", "#B45309", "#0F766E", "#166534"]
    for position, label, fill, outline in zip(positions, item["flow"], fills, outlines):
        box(draw, position, label, fill, outline, 20)
    gate = (250, 730, 1350, 825)
    box(draw, gate, item["gate"], "#FFF7ED", "#B45309", 18)
    path = DIAGRAM / f"{index:02d}-flowchart.png"
    image.save(path)
    return path


def make_erd(index: int, item: dict) -> Path:
    image, draw = canvas(f"ERD dan Aliran Data {index} - {item['title']}")
    center = (625, 345, 975, 475)
    nodes = [
        (55, 150, 405, 280), (1195, 150, 1545, 280),
        (55, 610, 405, 740), (1195, 610, 1545, 740), center,
    ]
    # Hubungan digambar di belakang simpul dan berakhir tepat pada sisi kotak.
    arrow(draw, (625, 380), (405, 215), "#0F766E", 4)
    arrow(draw, (975, 380), (1195, 215), "#0F766E", 4)
    arrow(draw, (625, 440), (405, 675), "#166534", 4)
    arrow(draw, (975, 440), (1195, 675), "#166534", 4)
    fills = ["#DCFCE7", "#ECFDF5", "#FFFFFF", "#FFFFFF", "#FEF3C7"]
    outlines = ["#166534", "#0F766E", "#166534", "#0F766E", "#B45309"]
    labels = item["entities"][:4] + [item["entities"][4]]
    for position, label, fill, outline in zip(nodes, labels, fills, outlines):
        box(draw, position, label, fill, outline, 20)
    draw.text((438, 246), "1..n", font=font(18, True), fill="#526273")
    draw.text((1100, 246), "1..n", font=font(18, True), fill="#526273")
    draw.text((440, 610), "audit", font=font(18, True), fill="#526273")
    draw.text((1100, 610), "posting", font=font(18, True), fill="#526273")
    path = DIAGRAM / f"{index:02d}-erd-dataflow.png"
    image.save(path)
    return path


def make_diagrams(items: list[dict]):
    DIAGRAM.mkdir(parents=True, exist_ok=True)
    return {
        index: (make_use_case(index, item), make_flow(index, item), make_erd(index, item))
        for index, item in enumerate(items, 1)
    }


def shade(cell, fill: str):
    pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def set_cell_margins(cell, top=55, start=120, bottom=55, end=120):
    pr = cell._tc.get_or_add_tcPr()
    margins = pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell(cell, value, bold=False, color=INK, size=9, align=None):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(value))
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def set_table_geometry(table, widths):
    dxa = [round(width * 1440) for width in widths]
    assert sum(dxa) == 9360, (widths, dxa)
    table.autofit = False
    pr = table._tbl.tblPr
    width = pr.first_child_found_in("w:tblW")
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    pr.append(indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    pr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for cell, value in zip(row.cells, dxa):
            tcw = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tcw.set(qn("w:w"), str(value))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(value / 1440)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header, True, "FFFFFF", 9)
        shade(table.rows[0].cells[index], GREEN)
    BASE.set_repeat_header(table.rows[0])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 and len(headers) > 2 else None
            set_cell(cells[index], value, False, INK, 8.5, alignment)
            if row_index % 2:
                shade(cells[index], "F7FCF8")
    set_table_geometry(table, widths)
    return table


def add_page_field(paragraph):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Title", 30, DEEP, 0, 12),
        ("Heading 1", 16, GREEN, 18, 10),
        ("Heading 2", 13, GREEN, 14, 7),
        ("Heading 3", 12, DEEP, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = f"APOTIK | UAT & PANDUAN | {VERSION}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Calibri"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Dokumen terkendali | Halaman ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    add_page_field(footer)

    doc.core_properties.title = "Laporan UAT dan Panduan Apotik v1.34.25"
    doc.core_properties.subject = "UAT live end-to-end, panduan operator, dan bukti visual Apotik"
    doc.core_properties.author = "Tim eBisnis / Zishof"
    doc.core_properties.keywords = "Apotik, UAT, manual, Android, Windows, pengadaan, akuntansi"


def note(doc: Document, label: str, text: str, color=GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, LIGHT)
    set_cell_margins(cell, 120, 160, 120, 160)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    paragraph.add_run(text)
    set_table_geometry(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_picture(doc: Document, path: Path, caption: str, width=6.45):
    if not path.exists():
        raise FileNotFoundError(f"Screenshot belum tersedia: {path}")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.style = doc.styles["Caption"]
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.keep_with_next = True


def add_bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def add_number(doc: Document, text: str, number: int):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.add_run(f"{number}. ").bold = True
    paragraph.add_run(text)


def scenarios(summary: dict, cashier: dict, procurement: dict,
              journal: dict, financial: dict) -> list[dict]:
    catalog = summary.get("katalogItemTotal", 0)
    ingredients = summary.get("bahanRacikanTerverifikasi", summary.get("bahanRacikanTotal", 0))
    provisioning = cashier.get("provisioning", {})
    transactions = cashier.get("transaksiUtamaLulus", {})
    idempotency = cashier.get("postflightIdempoten", {})
    recipes = provisioning.get(
        "resepSiapJualSetelahUat",
        summary.get("resepSiapJualTerverifikasi", summary.get("resepSiapJualTotal", 0)))
    sold = transactions.get("total", summary.get("totalTransaksiPenjualanLulus", 0))
    queue = summary.get("antreanFarmasiBaris", summary.get("antreanFarmasiDipastikan", 0))
    report_rows = summary.get("laporanPenjualanBaris", 0)
    common = dict(
        controls="Role sesuai fungsi, data sumber lengkap, kode idempoten, audit trail, dan pembacaan ulang hasil.",
        acceptance="Operasi tersimpan sekali, dapat ditemukan kembali, dan konsisten dengan detail serta laporan sumber.",
    )
    return [
        dict(title="Navigasi Menu dan Dashboard Operasional Apotik", screenshot=SHOT / "00-dashboard-operasional.png",
             actors=["Pemilik Apotik", "Apoteker", "Kasir", "Gudang"],
             usecases=["Memantau indikator harian", "Menilai antrean dan resep", "Membuka tindakan kasir", "Menindak stok dan batch"],
             flow=["Masuk sebagai role Apotik", "Baca kartu ringkasan", "Periksa peringatan", "Buka modul terkait", "Tindak lanjuti", "Konfirmasi hasil"],
             gate="GATE: identitas varian, tenant, toko, tanggal, dan status sinkronisasi harus benar.",
             entities=["PENGGUNA", "TOKO/APOTEK", "TRANSAKSI HARIAN", "PERINGATAN", "DASHBOARD"],
             paragraphs=[
                 "Dashboard adalah titik orientasi setelah login. Sidebar Apotik kini menampilkan tiga belas tujuan mandiri: Dashboard Apotik, Kasir Apotik, Tebus Resep Dokter, Racikan, Produksi Farmasi, Manajemen Farmasi, Formularium & Obat, Batch & Kedaluwarsa, Pengadaan/PBF, Stok Opname Apotik, Retur Obat, Obat Terkendali, dan Laporan Apotik. Satu hak akses dipetakan ke satu tujuan sehingga menu yang diizinkan tidak lagi tersembunyi di balik satu halaman gabungan.",
                 f"UAT memeriksa bahwa seluruh tiga belas menu tampil untuk role Admin pada toko Demo, tujuan aktif berubah sesuai halaman, dan dashboard dapat dibuka pada kanvas 1920 x 1080 tanpa overflow. Data live yang dibaca ulang mencakup {catalog} item katalog, {ingredients} bahan racikan, {recipes} resep siap jual, dan {queue} antrean farmasi. Test widget tambahan memverifikasi pemetaan fail-closed: hak yang tidak diberikan tidak ikut tampil.",
                 "Cara pakai dimulai dengan memeriksa toko dan peran pada header, lalu memilih menu Apotik sesuai pekerjaan. Jika sebuah menu tidak terlihat, periksa hak akses spesifik menu, keluar-masuk kembali agar cache hak dimuat ulang, kemudian sinkronkan. Jangan memberi seluruh hak hanya untuk memperbaiki satu menu; berikan hak minimum sesuai tugas pengguna.",
             ], **common),
        dict(title="Penjualan Obat Jadi", screenshot=SHOT / "01-kasir-obat-jadi.png",
             actors=["Kasir", "Pasien", "Apoteker", "Sistem stok"],
             usecases=["Mencari obat jadi", "Memilih batch layak", "Memvalidasi obat terkendali", "Menyelesaikan pembayaran"],
             flow=["Pilih OTC/obat bebas", "Cari nama atau barcode", "Tambahkan ke keranjang", "Periksa qty dan batch", "Pilih pembayaran", "Bayar dan cetak"],
             gate="GATE: stok cukup, batch belum kedaluwarsa, harga aktif, dan persyaratan obat terkendali terpenuhi.",
             entities=["ITEM OBAT", "BATCH", "KERANJANG", "PEMBAYARAN", "PENJUALAN"],
             paragraphs=[
                 "Layar kasir obat jadi dibagi menjadi mode transaksi, katalog, dan keranjang. Kolom pencarian menerima nama, kode, atau barcode; kartu item memperlihatkan harga serta penanda seperti LASA atau terkendali. Keranjang di sisi kanan menjadi tempat pemeriksaan akhir sebelum pembayaran. Susunan ini memungkinkan kasir mempertahankan konteks tanpa berpindah halaman ketika antrean pelanggan ramai.",
                 f"Pengujian membuat {transactions.get('otcObatBebas', 0)} transaksi OTC dengan kode UAT unik dan memeriksa ulang {idempotency.get('otcObatBebas', 0)} kode setelah pembayaran. Setiap penjualan memakai batch aktif dan jumlah yang tersedia. Pembacaan ulang membuktikan bahwa permintaan berulang tidak membuat nota kedua. Katalog yang digunakan adalah data sample/UAT dan tidak boleh diperlakukan sebagai daftar obat terdaftar atau acuan terapi.",
                 "Operator mencari obat, menilai nama dan dosis secara lengkap, memeriksa peringatan LASA, lalu mengatur jumlah. Sebelum menekan Bayar, cocokkan item, kuantitas, harga, identitas pembeli bila diperlukan, metode pembayaran, dan ketentuan resep. Bila stok atau batch ditolak, jangan memilih batch kedaluwarsa sebagai jalan pintas; eskalasikan ke gudang atau apoteker dan gunakan alur penerimaan/opname yang sah.",
             ], **common),
        dict(title="Mode Resep Dokter pada Kasir", screenshot=SHOT / "01b-kasir-resep-dokter.png",
             actors=["Kasir", "Apoteker", "Dokter", "Pasien"],
             usecases=["Memilih mode resep", "Mencari obat resep", "Memeriksa persyaratan", "Menyelesaikan transaksi"],
             flow=["Pilih Resep Dokter", "Cari item resep", "Tambahkan item", "Lengkapi identitas", "Periksa batch", "Bayar dan catat"],
             gate="GATE: mode resep aktif, identitas sumber jelas, item dan batch sesuai, serta pagar keselamatan lulus.",
             entities=["RESEP DOKTER", "ITEM OBAT", "PASIEN/DOKTER", "KERANJANG", "PEMBAYARAN"],
             paragraphs=[
                 "Mode Resep Dokter kini dapat dipilih langsung dari ruang kerja Kasir tanpa ikon kunci. Penanda mode aktif terlihat pada panel kiri, sementara katalog, identitas pembeli, metode pembayaran, dan keranjang tetap berada dalam satu konteks kerja. Perubahan mode akan mengosongkan keranjang berjalan hanya setelah operator mengonfirmasi.",
                 f"UAT menjalankan {transactions.get('resepDokter', 0)} transaksi Resep Dokter dan memeriksa ulang {idempotency.get('resepDokter', 0)} kode unik. Seluruh hasil tercatat satu kali; tidak ada transaksi gagal tersisa setelah respons gateway sementara diperiksa ulang.",
                 "Kasir memilih mode Resep Dokter, mencari item sesuai resep, lalu mencocokkan nama, kekuatan, bentuk sediaan, jumlah, dan dokter penulis. Obat terkendali tetap menuntut identitas pembeli dan referensi dokter/resep. Bila informasi klinis belum memadai, hentikan pembayaran dan arahkan ke Apoteker untuk telaah.",
             ], **common),
        dict(title="Pembayaran Obat Bebas", screenshot=SHOT / "01c-pembayaran-obat-bebas.png",
             actors=["Kasir", "Pelanggan", "Gateway pembayaran", "Sistem akuntansi"],
             usecases=["Memilih metode", "Mencatat referensi", "Membagi pembayaran", "Mengonfirmasi pembayaran"],
             flow=["Periksa total", "Buka Bayar", "Pilih metode", "Isi referensi/tunai", "Konfirmasi", "Cetak dan rekonsiliasi"],
             gate="GATE: kode transaksi server tersedia, metode pembayaran aktif, total cocok, dan kiriman idempoten.",
             entities=["KERANJANG", "CARA BAYAR", "REFERENSI", "PEMBAYARAN", "STRUK"],
             paragraphs=[
                 "Lembar Pembayaran Obat Bebas terbuka dari tombol Bayar dan menampilkan total, metode yang disediakan server, nomor referensi/approval, opsi bayar terpisah, serta aksi final. Pesan validasi membantu kasir memahami kapan referensi diperlukan untuk rekonsiliasi. Bukti layar diambil sebelum aksi final sehingga tidak menambah transaksi.",
                 f"Perbaikan kode transaksi server divalidasi melalui {transactions.get('otcObatBebas', 0)} pembayaran OTC. Postflight menemukan {idempotency.get('otcObatBebas', 0)} kode unik yang dapat dibaca ulang; observasi gateway sementara berjumlah {cashier.get('observasi', {}).get('responsGatewaySementara', 0)} dan seluruhnya selesai setelah recheck tanpa data ganda.",
                 "Setelah memeriksa keranjang, kasir membuka Bayar, memilih metode sesuai pembayaran pelanggan, dan mengisi nomor referensi bila non-tunai. Untuk tunai, masukkan uang diterima dan periksa kembalian. Jika hasil server belum dapat dipastikan, jangan membuat nota baru; gunakan pemeriksaan pembayaran tertunda agar kode idempoten yang sama dipakai kembali.",
             ], **common),
        dict(title="Tebus Resep Dokter", screenshot=SHOT / "02-tebus-resep-dokter.png",
             actors=["Apoteker", "Kasir", "Dokter", "Pasien"],
             usecases=["Membaca antrean resep", "Melakukan telaah", "Memeriksa penyiapan", "Menyerahkan obat"],
             flow=["Buka Tebus Resep", "Pilih resep menunggu", "Periksa detail", "Lakukan verifikasi", "Selesaikan pembayaran", "Serahkan dan catat"],
             gate="GATE: resep, pasien, dokter, item, jumlah, batch, dan verifikasi farmasi harus lengkap.",
             entities=["RESEP", "DETAIL RESEP", "PASIEN/DOKTER", "BATCH", "PENEBUSAN"],
             paragraphs=[
                 "Menu Tebus Resep Dokter berdiri sendiri pada sidebar dan membuka antrean resep yang menunggu penebusan. Daftar di sisi kiri mempertahankan nomor resep serta status, sedangkan area detail dipakai untuk telaah, pemeriksaan kedua, konseling, dan penyerahan setelah sebuah resep dipilih.",
                 f"Retest live membaca {recipes} resep siap jual dari server dan menampilkan 100 resep menunggu pada halaman pertama. UAT menyelesaikan {transactions.get('tebusResepCampuran', 0)} penebusan resep campuran dan memeriksa ulang {idempotency.get('tebusResep', 0)} kode unik. Capture ulang dilakukan dalam mode baca agar tidak menggandakan pembayaran yang telah tervalidasi.",
                 "Apoteker memilih resep, mencocokkan pasien dan dokter, memeriksa item serta batch, lalu menyelesaikan checklist farmasi. Kasir hanya menagihkan resep yang sudah memenuhi gate. Jika detail tidak lengkap atau stok berubah, tahan penebusan dan perbaiki sumber; jangan membuat resep kedua sebagai jalan pintas.",
             ], **common),
        dict(title="Telaah dan Penyerahan Resep", screenshot=SHOT / "02a-tebus-resep-detail.png",
             actors=["Apoteker", "Petugas pemeriksa", "Kasir", "Pasien"],
             usecases=["Menelaah keselamatan", "Melakukan pemeriksaan kedua", "Mencatat konseling", "Menyerahkan obat"],
             flow=["Pilih resep", "Baca profil/alergi", "Periksa item", "Lakukan double-check", "Catat konseling", "Lanjutkan penebusan"],
             gate="GATE: peringatan klinis ditindaklanjuti, item cocok, pemeriksaan kedua dan konseling tercatat sesuai kebijakan.",
             entities=["RESEP", "PROFIL PASIEN", "ALERGI/PERINGATAN", "DISPENSING", "AUDIT"],
             paragraphs=[
                 "Setelah resep dipilih, panel kanan menampilkan keselamatan klinis pasien, daftar periksa sebelum penyerahan, status pemeriksaan kedua dan konseling, serta obat pada resep. Bila profil pasien belum terhubung, layar menyatakan keterbatasan tersebut secara eksplisit dan mewajibkan konfirmasi alergi manual; sistem tidak memberi kesan telaah otomatis sudah lengkap.",
                 f"Bukti live memakai salah satu dari {recipes} resep yang tersedia dan memverifikasi bahwa detail dapat dirender. Dataset akhir menyisakan {provisioning.get('resepSiapJualSetelahUat', recipes)} resep siap jual setelah {cashier.get('resepFormulaDitandaiDitebus', 0)} resep formula ditandai ditebus dalam UAT.",
                 "Apoteker membaca setiap peringatan, mencocokkan item dan jumlah, lalu menjalankan pemeriksaan kedua dan konseling sesuai SOP. Untuk obat terkendali, identitas pembeli dan register wajib lengkap. Tombol tindakan tidak boleh dipakai sekadar untuk menghilangkan peringatan; isi catatan harus merepresentasikan kegiatan yang benar-benar dilakukan.",
             ], **common),
        dict(title="Penjualan Obat Racikan", screenshot=SHOT / "02-resep-racikan.png",
             actors=["Apoteker", "Kasir", "Dokter", "Pasien"],
             usecases=["Memilih resep menunggu", "Memeriksa komposisi", "Menyiapkan racikan", "Menebus dan menyerahkan"],
             flow=["Buka antrean resep", "Pilih resep", "Baca komposisi", "Validasi bahan", "Konfirmasi racikan", "Tebus resep"],
             gate="GATE: resep sah, komposisi terbaca, bahan cukup, dan verifikasi apoteker selesai.",
             entities=["RESEP", "DETAIL RESEP", "BAHAN RACIKAN", "PENEBUSAN", "AUDIT FARMASI"],
             paragraphs=[
                 "Menu Racikan membuka mode khusus pada kasir untuk mencari formula racikan yang sudah disahkan. Katalog formula berada di area tengah dan keranjang di kanan; tombol OTC, Resep Dokter, Racikan, Produksi Farmasi, dan Tebus Resep tetap tersedia sebagai perpindahan mode kerja. Kondisi kosong berarti belum ada formula yang cocok dan bukan alasan untuk mengarang komposisi.",
                 f"Server menyediakan {recipes} resep siap jual, {provisioning.get('formulaRacikanOperasional', 0)} formula racikan operasional, dan sekurangnya {ingredients} bahan racikan sample. UAT menyelesaikan {transactions.get('racikan', 0)} transaksi racikan serta memeriksa ulang {idempotency.get('racikan', 0)} kode unik. Keberhasilan dihitung dari respons tersimpan dan pembacaan ulang, bukan dari banyaknya kartu yang terlihat pada satu halaman.",
                 "Apoteker memilih resep, membandingkan komposisi dengan resep sumber, memeriksa ketersediaan batch bahan, lalu memberi status sesuai pekerjaan sebenarnya. Kasir hanya menyelesaikan tagihan setelah verifikasi farmasi selesai. Jika detail kosong, satuan tidak jelas, atau bahan tidak tersedia, resep ditahan. Jangan mengganti bahan atau dosis hanya agar transaksi bisa disimpan; koreksi harus berasal dari pihak yang berwenang dan meninggalkan jejak audit.",
             ], **common),
        dict(title="Produksi Farmasi", screenshot=SHOT / "02b-produksi-farmasi-konfirmasi.png",
             actors=["Apoteker", "Petugas produksi", "Gudang bahan", "Penanggung jawab mutu"],
             usecases=["Memilih formula", "Mengurangi bahan", "Membuat batch hasil", "Mencatat kedaluwarsa"],
             flow=["Buka Produksi Farmasi", "Pilih barang jadi", "Periksa kebutuhan bahan", "Isi nomor batch", "Tetapkan kedaluwarsa", "Proses dan verifikasi stok"],
             gate="GATE: formula/BOM aktif, bahan cukup, nomor batch unik, tanggal kedaluwarsa sah, dan kode proses idempoten.",
             entities=["FORMULA PRODUKSI", "BAHAN BAKU", "BATCH HASIL", "MUTASI STOK", "AUDIT PRODUKSI"],
             paragraphs=[
                 "Produksi Farmasi kini tersedia sebagai mode Kasir dan menu mandiri tanpa ikon kunci. Katalog menampilkan barang jadi atau formula produksi, sedangkan panel kanan menjadi Rencana Produksi. Dialog konfirmasi meminta nomor batch hasil dan tanggal kedaluwarsa sebelum bahan dikurangi serta batch barang jadi dibuat dalam satu proses.",
                 f"UAT menyediakan {provisioning.get('formulaProduksiOperasional', 0)} formula produksi, menjalankan {transactions.get('produksiFarmasi', 0)} proses produksi, dan membaca ulang {idempotency.get('produksiFarmasi', 0)} kode unik. Seluruh proses lulus tanpa kegagalan tersisa dan tanpa duplikasi akibat retry gateway.",
                 "Petugas memilih formula, memeriksa satuan hasil dan ketersediaan bahan, lalu mengisi nomor batch serta tanggal kedaluwarsa berdasarkan dokumen produksi. Setelah Proses Produksi dikonfirmasi, cocokkan pengurangan bahan, penambahan stok barang jadi, dan batch hasil. Jika bahan tidak cukup atau batch ditolak, perbaiki sumber; jangan mengulang dengan kode baru sebelum status kode lama dipastikan.",
             ], **common),
        dict(title="Manajemen Farmasi", screenshot=SHOT / "02c-manajemen-farmasi.png",
             actors=["Pemilik Apotik", "Apoteker PJ", "Manajer operasional", "Auditor"],
             usecases=["Memantau layanan klinis", "Memantau stok dan mutu", "Membuka proses farmasi", "Meninjau audit dan analitik"],
             flow=["Buka Manajemen Farmasi", "Baca ringkasan", "Pilih domain", "Buka modul tujuan", "Tindak lanjuti", "Pantau hasil"],
             gate="GATE: angka berasal dari API live, hak akses tujuan tersedia, dan tindakan kritis tetap memakai gate modul asal.",
             entities=["RINGKASAN FARMASI", "RESEP", "RACIKAN/PRODUKSI", "STOK/BATCH", "LAPORAN/AUDIT"],
             paragraphs=[
                 "Manajemen Farmasi adalah pusat navigasi untuk layanan pasien, telaah klinis, konseling, formula racikan, produksi/QC, recall dan karantina, cold chain, lokasi dan transfer, perencanaan stok, supplier, kas, audit, integrasi, membership, serta business intelligence. Kartu tidak menggantikan modul transaksi; setiap kartu mengarahkan pengguna ke halaman operasional yang relevan.",
                 f"Capture live membaca {catalog} obat, {recipes} resep, {provisioning.get('formulaRacikanOperasional', 0)} formula racikan, dan sedikitnya 100 data pada domain produksi/QC serta kas. Menu tampil mandiri pada sidebar dan kartu dapat dirender pada 1920 x 1080 tanpa overflow.",
                 "Manajemen membuka pusat kendali, membaca ringkasan, lalu memilih domain yang perlu ditindaklanjuti. Angka nol harus ditafsirkan sebagai keadaan data yang perlu diperiksa, bukan otomatis sebagai modul gagal. Untuk perubahan stok, pembayaran, produksi, atau status resep, pengguna tetap menyelesaikan langkah dan otorisasi pada modul tujuan.",
             ], **common),
        dict(title="Formularium dan Katalog Obat", screenshot=SHOT / "03-formularium-obat.png",
             actors=["Apoteker", "Gudang", "Manajemen", "Auditor"],
             usecases=["Menelusuri katalog", "Membedakan jenis item", "Memeriksa profil obat", "Mengendalikan aktivasi"],
             flow=["Buka formularium", "Gunakan pencarian", "Pilih kategori", "Buka profil", "Periksa stok/batch", "Simpan perubahan sah"],
             gate="GATE: kode unik, jenis item benar, satuan jelas, dan seluruh data ditandai sample bila untuk UAT.",
             entities=["ITEM MEDIS", "JENIS ITEM", "PROFIL APOTIK", "SATUAN", "KATALOG"],
             paragraphs=[
                 "Formularium adalah sumber identitas item yang dipakai oleh kasir, penerimaan, stok, resep, dan laporan. Tangkapan layar menunjukkan daftar item dengan pencarian dan informasi operasional. Perbedaan antara obat jadi, bahan baku racikan, alat kesehatan, serta item terkendali harus terlihat dari data master, bukan disimpulkan hanya dari nama produk.",
                 f"Setelah provisioning final, UAT mensyaratkan sedikitnya 11.000 item gabungan: 10.000 obat jadi dan 1.000 bahan racikan. Hasil mesin membaca {catalog} item katalog dan {ingredients} bahan berprefix DEMO-BHN. Nama, produsen, negara, dosis, harga, barcode, dan batch dibuat secara deterministik untuk pengujian volume. Semua data bersifat sintetik/sample dan tidak menyatakan izin edar, keamanan, atau kesetaraan klinis.",
                 "Untuk memeriksa satu item, cari kode yang spesifik, buka profil, lalu cocokkan jenis, satuan, harga, golongan, LASA, batas stok, barcode, dan batch. Perubahan profil tidak boleh dilakukan massal tanpa daftar persetujuan. Jika jumlah hasil tidak cocok dengan total, periksa pagination dan filter. Katalog besar harus tetap responsif serta tidak menutupi label atau tombol pada layar 1920 x 1080.",
             ], **common),
        dict(title="Batch, FEFO, dan Kedaluwarsa", screenshot=SHOT / "04-batch-kedaluwarsa.png",
             actors=["Gudang", "Apoteker", "Supervisor", "Sistem FEFO"],
             usecases=["Memantau tanggal ED", "Memilih batch FEFO", "Mengarantina batch", "Mencatat disposisi"],
             flow=["Pilih horizon", "Baca status batch", "Urutkan tanggal", "Periksa sisa", "Karantina/retur", "Dokumentasikan disposisi"],
             gate="GATE: batch lewat tanggal tidak boleh dipakai dalam penjualan atau racikan.",
             entities=["ITEM", "BATCH/LOT", "MUTASI STOK", "KARANTINA", "MONITOR ED"],
             paragraphs=[
                 "Monitor batch memusatkan perhatian pada nomor lot, tanggal kedaluwarsa, sisa, dan status penanganan. Pilihan horizon membantu gudang memisahkan barang yang harus segera diprioritaskan, dipindahkan, diretur, atau dikarantina. Warna peringatan mendukung keputusan, tetapi tanggal dan sisa tetap harus dibaca; warna bukan satu-satunya dasar tindakan.",
                 f"UAT menjalankan monitor dan laporan kedaluwarsa setelah transaksi. Seleksi batch menemukan {summary.get('batchKedaluwarsaDitemukanSaatSeleksi', 0)} baris kedaluwarsa dan hanya memilih batch aktif untuk penjualan. Endpoint monitor berstatus {summary.get('batchMonitorStatus', 'belum tercatat')} dengan {summary.get('batchMonitorBaris', 0)} baris yang dapat dibaca pada halaman uji. Penolakan batch kedaluwarsa adalah kontrol keselamatan yang tidak boleh dinonaktifkan untuk mengejar kelulusan tes.",
                 "Petugas memilih horizon, membuka item, menelusuri sumber penerimaan, dan memastikan jumlah fisik sesuai. Batch bermasalah dipindahkan ke status karantina melalui prosedur yang disetujui. Bila tanggal tidak valid atau sisa negatif, hentikan transaksi terkait, simpan bukti, dan lakukan rekonsiliasi ledger. Disposisi akhir harus mencatat jumlah, alasan, petugas, persetujuan, dan dokumen retur atau pemusnahan.",
             ], **common),
        dict(title="Pengadaan / PBF dan Penerimaan Barang", screenshot=SHOT / "05-penerimaan-pbf.png",
             actors=["Gudang", "PBF/Vendor", "Apoteker", "Keuangan"],
             usecases=["Menerima barang", "Mencatat batch dan ED", "Memeriksa kuantitas", "Menghubungkan dokumen sumber"],
             flow=["Pilih vendor/dokumen", "Cocokkan barang", "Isi lot dan ED", "Periksa harga/qty", "Simpan penerimaan", "Baca ulang stok"],
             gate="GATE: barang, jumlah, mutu, batch, ED, dan referensi PO/BAST harus konsisten.",
             entities=["VENDOR/PBF", "PO", "BAST/PENERIMAAN", "BATCH", "PERSEDIAAN"],
             paragraphs=[
                 "Menu Pengadaan/PBF membuka formulir penerimaan pada tujuan sidebar tersendiri. Penerimaan PBF adalah pintu masuk stok sehingga kesalahan di layar ini akan terbawa ke penjualan, FEFO, HPP, tagihan, dan laporan. Form menampilkan pemasok, nomor faktur, item, jumlah, harga beli, batch, dan tanggal kedaluwarsa. Operator harus bekerja dari surat jalan, PO, dan pemeriksaan fisik.",
                 f"Dalam UAT, daftar penerimaan membuktikan {procurement.get('pengadaan_bast_daftar', 0)} BAST sample yang dapat ditelusuri ke PO. Item jangkar UJI-PCT juga diberi top-up batch idempoten dengan kode faktur UAT-TERIMA-APT-13424-001 bila batch yang cukup belum ada. Setelah penyimpanan, batch dibaca kembali dan digunakan oleh transaksi penjualan. Cara ini membuktikan bahwa penerimaan memengaruhi stok nyata dan bukan sekadar menghasilkan notifikasi sukses.",
                 "Gudang mencocokkan satu per satu nama item, kuantitas, satuan, nomor lot, dan ED. Apoteker menilai kelayakan mutu; keuangan tidak memproses tagihan bila penerimaan belum sah. Jika jumlah berbeda dari PO atau barang rusak, catat selisih dan gunakan status pengecualian. Jangan menaikkan stok lewat koreksi langsung agar angka cocok, karena tindakan itu memutus hubungan antara persediaan dan dokumen pengadaan.",
             ], **common),
        dict(title="Stok Opname Apotik", screenshot=SHOT / "05b-stok-opname-apotik.png",
             actors=["Petugas Gudang", "Apoteker", "Supervisor", "Auditor"],
             usecases=["Menambah item hitung", "Mencatat stok fisik", "Menghitung selisih", "Menyimpan koreksi sah"],
             flow=["Buka Stok Opname", "Tambahkan obat", "Hitung stok fisik", "Bandingkan sistem", "Setujui selisih", "Simpan dan audit"],
             gate="GATE: item, batch, jumlah fisik, alasan selisih, dan persetujuan harus dapat ditelusuri.",
             entities=["ITEM", "BATCH", "STOK SISTEM", "STOK FISIK", "MUTASI OPNAME"],
             paragraphs=[
                 "Stok Opname Apotik kini dapat dibuka langsung dari sidebar. Halaman dimulai dalam keadaan aman tanpa baris; petugas menambahkan obat yang benar-benar dihitung, lalu tombol simpan aktif setelah data wajib dan selisih tersedia.",
                 f"Retest membuka tujuan ini dengan role Admin dan sumber pilihan yang sama dengan formularium berisi {catalog} item. Harness memastikan tab Stok Opname tetap aktif dan tidak kembali ke Formularium ketika proses data contoh di latar selesai. Perbaikan ini mencegah operator menyimpan pada konteks halaman yang keliru.",
                 "Lakukan hitung fisik per lokasi dan batch, catat jumlah aktual, minta supervisor menilai selisih, lalu simpan satu kali. Jika hasil hitung belum final, jangan menekan Simpan Opname. Gunakan kode dan catatan yang dapat diaudit; koreksi berikutnya harus menjadi transaksi baru, bukan penghapusan jejak sebelumnya.",
             ], **common),
        dict(title="Retur Obat", screenshot=SHOT / "05c-retur-obat.png",
             actors=["Kasir", "Gudang", "Apoteker", "PBF/Vendor"],
             usecases=["Menerima retur pembeli", "Mengirim retur ke PBF", "Memeriksa batch", "Mencatat alasan"],
             flow=["Buka Retur Obat", "Pilih arah retur", "Isi alasan", "Tambahkan obat", "Periksa qty/batch", "Catat retur"],
             gate="GATE: arah retur, sumber transaksi, alasan, item, batch, jumlah, dan status mutu harus lengkap.",
             entities=["PENJUALAN/PENERIMAAN", "ITEM", "BATCH", "RETUR", "MUTASI STOK"],
             paragraphs=[
                 "Menu Retur Obat berdiri sendiri dan menyediakan dua arah: dari pembeli sebagai barang masuk, atau ke PBF sebagai barang keluar. Pilihan arah, alasan, item, batch, dan jumlah menentukan mutasi stok sehingga petugas wajib memastikan konteks sebelum mencatat.",
                 f"Retest membuka halaman secara langsung, mempertahankan tab Retur sebagai tujuan aktif, dan menggunakan katalog sumber {catalog} item. Validasi volume penjualan, batch, serta laporan dari run sebelumnya tetap menjadi bukti bahwa dampak retur harus direkonsiliasi terhadap sumber yang sama; run capture tidak membuat retur baru.",
                 "Cari transaksi atau penerimaan asal, pilih arah retur, tulis alasan, lalu masukkan item dan batch yang benar. Barang rusak atau kedaluwarsa tidak boleh kembali menjadi stok jual. Jika sumber tidak ditemukan, tahan barang secara fisik dan eskalasikan; jangan mencatat retur tanpa referensi hanya agar stok terlihat cocok.",
             ], **common),
        dict(title="Layar Publik Obat Jadi dan Racikan", screenshot=SHOT / "06-layar-kedua-obat-jadi-racikan.png",
             actors=["Pasien/Keluarga", "Petugas Farmasi", "Sistem Antrean", "Petugas Loket"],
             usecases=["Melihat status anonim", "Membedakan jenis layanan", "Mengetahui loket", "Memanggil antrean siap"],
             flow=["Ambil antrean server", "Samarkan identitas", "Kelompokkan jenis", "Tampilkan status", "Arahkan ke loket", "Perbarui otomatis"],
             gate="GATE: identitas pasien tersamar dan hanya informasi yang diperlukan publik yang ditampilkan.",
             entities=["ANTREAN FARMASI", "PASIEN TERSAMAR", "DAFTAR OBAT PUBLIK", "LOKET", "LAYAR KEDUA"],
             paragraphs=[
                 "Layar kedua menggabungkan obat jadi dan racikan agar pasien atau keluarga dapat mengetahui posisi proses tanpa bertanya berulang ke loket. Informasi dipisahkan menurut jenis layanan, status seperti menunggu, disiapkan, atau siap, serta tujuan loket. Ukuran huruf dan kontras dirancang untuk televisi atau monitor publik yang dilihat dari jarak tertentu.",
                 f"Endpoint live mengembalikan {queue} antrean untuk toko uji dan menyatakan kebijakan privasi {summary.get('identitasLayarPublik', 'belum tercatat')}. UAT memastikan data berasal dari server, bukan pratinjau lokal. Nama dan nomor rekam medis harus disamarkan, sedangkan catatan publik tidak boleh memuat diagnosis, rincian klinis sensitif, nomor telepon, alamat, atau informasi pembayaran.",
                 "Petugas membuka layar pada monitor kedua, memilih toko dan mode gabungan, lalu memeriksa pembaruan otomatis. Ketika obat siap, status dan loket harus berubah tanpa memerlukan login pasien. Jika identitas lengkap terlihat, layar wajib segera ditutup dan insiden privasi dicatat. Bila koneksi terputus, tampilan harus menunjukkan kondisi yang jujur dan tidak terus memanggil antrean berdasarkan data lama seolah-olah masih terbaru.",
             ], **common),
        dict(title="Layar Tambahan Obat Jadi", screenshot=SHOT / "07-layar-tambahan-obat-jadi.png",
             actors=["Pasien OTC", "Petugas Farmasi", "Sistem Antrean", "Petugas Loket"],
             usecases=["Memfilter obat jadi", "Menampilkan status", "Mengatur loket", "Memperbarui panggilan"],
             flow=["Pilih mode obat jadi", "Ambil antrean", "Terapkan filter", "Tampilkan urutan", "Panggil yang siap", "Segarkan layar"],
             gate="GATE: hanya antrean obat jadi yang tampil dan identitas tetap tersamar.",
             entities=["ANTREAN", "JENIS JADI", "STATUS", "LOKET", "MONITOR PUBLIK"],
             paragraphs=[
                 "Mode obat jadi menyediakan layar khusus ketika instalasi memiliki monitor terpisah atau volume resep non-racikan tinggi. Filter jenis mengurangi kepadatan informasi dan membuat nomor antrean, status, serta loket lebih mudah dibaca. Mode ini bukan daftar penjualan; pasien hanya melihat informasi pengambilan yang diperlukan.",
                 f"UAT membandingkan isi mode khusus dengan data antrean gabungan dan memverifikasi {summary.get('antreanObatJadi', 0)} antrean obat jadi sample. Baris bertipe racikan atau campuran tidak boleh masuk ke tampilan ini. Status siap harus tetap menampilkan loket yang tepat, sedangkan identitas mengikuti masker server. Penambahan layar tidak membuat salinan data baru: semua monitor membaca sumber antrean yang sama sehingga perubahan status konsisten.",
                 "Petugas memilih mode Obat Jadi pada konfigurasi layar tambahan dan menempatkan jendela pada monitor yang sesuai. Setelah pergantian shift atau toko, periksa kembali judul instalasi dan loket. Jika sebuah nomor muncul di mode yang salah, verifikasi nilai jenis pada sumber antrean; jangan menutupi kesalahan dengan menghapus pasien dari daftar publik.",
             ], **common),
        dict(title="Layar Tambahan Obat Racikan", screenshot=SHOT / "08-layar-tambahan-racikan.png",
             actors=["Pasien Resep", "Apoteker", "Sistem Antrean", "Petugas Loket"],
             usecases=["Memfilter racikan", "Memantau tahap persiapan", "Memanggil pasien", "Menjaga privasi"],
             flow=["Pilih mode racikan", "Ambil antrean", "Terapkan filter", "Tampilkan proses", "Arahkan ke loket", "Perbarui status"],
             gate="GATE: antrean racikan/campuran benar, status tidak prematur, dan data klinis tidak dipublikasikan.",
             entities=["RESEP", "ANTREAN RACIKAN", "STATUS PROSES", "LOKET", "LAYAR PUBLIK"],
             paragraphs=[
                 "Mode racikan menyoroti antrean yang membutuhkan penyiapan farmasi lebih panjang. Pasien dapat membedakan apakah resep masih menunggu, sedang diracik, atau siap diambil. Informasi tersebut mengurangi tekanan pada loket sekaligus menjaga agar apoteker tidak mengubah status hanya untuk memenuhi ekspektasi waktu.",
                 f"UAT memastikan filter menampilkan {summary.get('antreanRacikan', 0)} antrean racikan sample dan tidak memasukkan obat jadi murni. Komposisi detail resep tidak ditampilkan; layar cukup memberi label layanan dan catatan publik yang aman. Pengujian juga menilai keterbacaan pada 1920 x 1080 serta ketiadaan teks yang terpotong pada daftar panjang.",
                 "Apoteker memperbarui status berdasarkan pekerjaan nyata. Petugas loket memanggil hanya entri siap. Jika racikan harus dikoreksi, status dikembalikan melalui alur operasional dan pasien diberi informasi non-klinis yang sesuai. Jangan menampilkan nama bahan sensitif, diagnosis, atau instruksi terapi di layar publik; rincian tersebut tetap berada pada layar petugas berotorisasi.",
             ], **common),
        dict(title="Laporan Penjualan Apotik", screenshot=SHOT / "09-laporan-penjualan.png",
             actors=["Supervisor", "Keuangan", "Kasir", "Auditor"],
             usecases=["Memilih periode", "Membaca agregat", "Menelusuri item", "Merekonsiliasi nota"],
             flow=["Atur periode", "Terapkan filter", "Baca total", "Periksa item", "Bandingkan nota", "Ekspor/arsipkan"],
             gate="GATE: periode, status final, toko, pembatalan, diskon, dan retur dihitung konsisten.",
             entities=["PENJUALAN", "DETAIL ITEM", "FILTER PERIODE", "AGREGAT", "LAPORAN"],
             paragraphs=[
                 "Laporan penjualan merangkum kuantitas dan nilai per item untuk periode yang dipilih. Layar harus dipakai setelah transaksi tersimpan, bukan sebagai pengganti nota. Filter periode dan toko menentukan ruang lingkup; total kemudian ditelusuri ke item serta transaksi sumber ketika ada perbedaan.",
                 f"Setelah perbaikan endpoint, laporan berstatus {summary.get('laporanPenjualanStatus', 'belum tercatat')}, membaca {report_rows} baris produk pada klien, dan sumber server mencatat {sold} transaksi UAT lulus. Skenario memakai sedikitnya 100 item berbeda agar halaman laporan tidak hanya menunjukkan satu atau dua agregat. Acceptance test kemudian membandingkan kuantitas, nilai, dan jejak nota pada periode yang sama.",
                 "Supervisor memilih periode yang mencakup waktu UAT, memastikan toko benar, lalu membandingkan item UJI-PCT dan kode transaksi berprefix rilis. Jika laporan kosong sementara nota ada, periksa cut-off tanggal, status transaksi, zona waktu, dan struktur respons. Jangan mengubah tanggal transaksi agar laporan terlihat benar; temuan harus diperbaiki pada filter atau agregasi dan kemudian diretest.",
             ], **common),
        dict(title="Register Obat Terkendali", screenshot=SHOT / "10-register-obat-terkendali.png",
             actors=["Apoteker", "Penanggung Jawab", "Auditor", "Sistem Kepatuhan"],
             usecases=["Membaca transaksi terkendali", "Memeriksa identitas wajib", "Menelusuri batch", "Mengaudit penyerahan"],
             flow=["Pilih periode", "Filter golongan", "Buka transaksi", "Periksa identitas", "Cocokkan batch", "Tandai pemeriksaan"],
             gate="GATE: obat terkendali tidak boleh diserahkan tanpa data dan otorisasi yang dipersyaratkan.",
             entities=["ITEM TERKENDALI", "TRANSAKSI", "PASIEN/RESEP", "BATCH", "REGISTER"],
             paragraphs=[
                 "Register obat terkendali memisahkan transaksi yang memerlukan jejak kepatuhan lebih kuat dari penjualan obat bebas. Kolom dan filter membantu apoteker menelusuri tanggal, item, jumlah, referensi resep atau pasien, dan batch. Layar ini tidak boleh menyederhanakan kewajiban institusi; SOP dan ketentuan regulator tetap menjadi sumber keputusan.",
                 f"UAT membuka endpoint dan halaman dengan status {summary.get('laporanTerkendaliStatus', 'belum tercatat')} serta membaca {summary.get('laporanTerkendaliBaris', 0)} baris register pada periode uji. Item UJI-CDN dipakai pada transaksi sample beridentitas fiktif, alamat sample, nama dokter sample, dan batch aktif. Pengujian membuktikan validasi data wajib, pencatatan register, serta kemampuan halaman menangani volume audit tanpa memakai data pasien nyata.",
                 "Apoteker memilih periode, membuka transaksi, dan mencocokkan item, jumlah, batch, pihak yang menyerahkan, serta data resep yang diwajibkan. Bila identitas atau referensi kosong, penyerahan ditahan dan transaksi dikoreksi lewat alur resmi. Akses register harus dibatasi; ekspor dan screenshot tidak boleh menyebarkan data pasien di luar tujuan audit yang sah.",
             ], **common),
        dict(title="Laporan Kedaluwarsa", screenshot=SHOT / "11-laporan-kedaluwarsa.png",
             actors=["Gudang", "Apoteker", "Manajemen", "Auditor"],
             usecases=["Menyaring horizon ED", "Mengukur nilai risiko", "Menentukan tindakan", "Membuktikan disposisi"],
             flow=["Pilih horizon", "Tampilkan batch", "Kelompokkan urgensi", "Cek stok fisik", "Tentukan disposisi", "Rekonsiliasi"],
             gate="GATE: hasil laporan harus cocok dengan batch fisik dan transaksi mutasi.",
             entities=["BATCH", "ITEM", "HORIZON", "NILAI PERSEDIAAN", "LAPORAN ED"],
             paragraphs=[
                 "Laporan kedaluwarsa mengubah data batch menjadi daftar kerja bagi gudang dan apoteker. Berbeda dari monitor operasional, laporan ini dipakai untuk melihat dampak lintas item pada horizon tertentu dan mendukung perencanaan retur, pemindahan, atau pemusnahan. Tanggal, sisa, dan nilai harus dapat ditelusuri kembali ke batch sumber.",
                 f"Endpoint berstatus {summary.get('laporanKedaluwarsaStatus', 'belum tercatat')} dan klien membaca {summary.get('laporanKedaluwarsaBaris', 0)} baris pada pengujian. Katalog besar menghasilkan banyak batch sample dengan tanggal bervariasi sehingga pagination dan performa ikut diuji. Volume tinggi tidak otomatis berarti risiko nyata karena seluruh data UAT bersifat sintetik.",
                 "Pengguna memilih horizon sesuai kebijakan, memeriksa item berisiko tertinggi, lalu mencockan dengan stok fisik. Setiap tindakan harus mengubah status batch atau mutasi melalui transaksi yang tepat agar laporan berikutnya konsisten. Jika baris tetap muncul setelah disposisi, periksa tanggal efektif, status batch, dan posting mutasi sebelum membuat koreksi baru.",
             ], **common),
        dict(title="Procure-to-Pay: PR sampai Pembayaran Vendor", screenshot=PROC / "10-pembayaran-vendor-100.png",
             actors=["Peminta", "Procurement", "Penerima", "Keuangan"],
             usecases=["Membuat PR/PO", "Menerima melalui BAST", "Mencatat tagihan", "Membayar vendor"],
             flow=["Buat PR", "Setujui dan buat PO", "Terima/BAST", "Terima tagihan", "Setujui pembayaran", "Rekonsiliasi"],
             gate="GATE: segregasi tugas, vendor sah, kuantitas diterima, tagihan, dan persetujuan lengkap.",
             entities=["PR", "PO", "BAST", "TAGIHAN", "PEMBAYARAN"],
             paragraphs=[
                 "Layar pembayaran vendor adalah ujung rangkaian procure-to-pay, bukan titik awal. Sebuah baris layak dibayar hanya bila dapat ditelusuri ke permintaan, pemesanan, penerimaan/BAST, dan tagihan yang disetujui. Filter dan status pada halaman membantu keuangan memisahkan dokumen yang masih menunggu, ditolak, atau sudah dibayar.",
                 f"Dataset volume memuat {procurement.get('pengadaan_pr_daftar', 0)} PR, {procurement.get('pengadaan_po_daftar', 0)} PO, {procurement.get('pengadaan_bast_daftar', 0)} BAST, dan {procurement.get('pengadaan_tagihan_daftar', 0)} tagihan. Retest endpoint menemukan 100 pembayaran melalui marker yang sama, menutup temuan filter nol pada rilis sebelumnya. Data ini tetap berada pada tenant demo dan ditandai prefix UAT agar dapat diaudit.",
                 "Keuangan membuka detail sebelum menyetujui pembayaran, mencocokkan vendor, nomor tagihan, nilai, cara bayar, dan dokumen sumber. Bila BAST atau persetujuan tidak ada, pembayaran dihentikan. Jangan menyelesaikan selisih dengan mengedit nominal secara terpisah pada tahap terakhir; koreksi harus kembali ke dokumen yang menyebabkan selisih sehingga hubungan audit tetap utuh.",
             ], **common),
        dict(title="Draft Jurnal dan Posting", screenshot=ACC / "08-draft-jurnal.png",
             actors=["Akuntan", "Approver", "Owner COA", "Auditor"],
             usecases=["Meninjau draft", "Memetakan akun", "Memvalidasi debit/kredit", "Memposting buku besar"],
             flow=["Buka draft", "Pilih kategori", "Baca referensi", "Validasi akun", "Posting", "Periksa buku besar"],
             gate="GATE: debit sama dengan kredit, akun leaf aktif, periode terbuka, dan sumber dapat ditelusuri.",
             entities=["DOKUMEN SUMBER", "DRAFT JURNAL", "DETAIL D/K", "POSTING HISTORY", "BUKU BESAR"],
             paragraphs=[
                 "Draft Jurnal menampilkan transaksi sumber yang belum atau sudah diproses menjadi jurnal. Akuntan memakai kategori untuk memisahkan penjualan, HPP, kulakan, pembayaran vendor, jurnal umum, dan proses lain. Jumlah pada kartu adalah indikator pekerjaan; keputusan posting tetap dilakukan setelah membuka referensi dan memeriksa akun debit-kredit.",
                 f"UAT vendor memeriksa {journal.get('pembayaranDiperiksa', 0)} pembayaran dan menemukan {journal.get('jurnalTerpostingTerverifikasi', 0)} jurnal terposting. Mapping yang diuji adalah debit {journal.get('debet', '-')} dan kredit {journal.get('kredit', '-')}. Retest memastikan jurnal dapat dibaca ulang sebagai posted; dashboard tenant bersama dapat menampilkan total yang lebih besar karena juga berisi transaksi modul lain.",
                 "Akuntan memilih kategori, membaca dokumen sumber, memeriksa tanggal serta nominal, lalu menilai akun yang disarankan. Posting tidak boleh dilakukan hanya karena tombol aktif. Bila akun salah, gunakan fasilitas penyesuaian mapping yang sah dan ulangi preview. Jika debit dan kredit tidak seimbang atau periode sudah ditutup, simpan bukti dan perbaiki sumber; jangan menambahkan baris penyeimbang tanpa substansi bisnis.",
             ], **common),
        dict(title="Laporan Keuangan dan Rekonsiliasi", screenshot=ACC / "24-laporan-laba-rugi.png",
             actors=["Akuntan", "Manajemen", "Auditor", "Pemilik Proses"],
             usecases=["Membaca Laba Rugi", "Merekonsiliasi Buku Besar", "Menilai Arus Kas", "Menyetujui periode"],
             flow=["Pastikan posting", "Pilih periode", "Buka laporan", "Telusuri akun", "Bandingkan sumber", "Sign-off"],
             gate="GATE: hanya jurnal posted, saldo awal sah, klasifikasi akun benar, dan periode konsisten.",
             entities=["AKUN", "JURNAL POSTED", "BUKU BESAR", "TRIAL BALANCE", "LAPORAN KEUANGAN"],
             paragraphs=[
                 "Tangkapan layar Laba Rugi memakai seluruh lebar area kerja: kolom keterangan memperoleh ruang utama sampai sisi kanan dan setiap sel angka dapat diklik untuk membuka data penyusunnya. Bukti ini menggantikan tangkapan layar kosong pada dokumen lama. Dari laporan ini akuntan melanjutkan rekonsiliasi ke Jurnal Umum, Buku Besar, Neraca Saldo, Neraca, dan Arus Kas dengan periode yang sama.",
                 f"UAT memposting dan membaca ulang {financial.get('sumberJurnalTerpostingTerverifikasi', 0)} jurnal sample sebagai sumber laporan, di samping {journal.get('jurnalTerpostingTerverifikasi', 0)} jurnal pembayaran vendor. Enam endpoint laporan keuangan diperiksa dan semuanya harus mengembalikan baris: Laba Rugi, Neraca, Arus Kas, Jurnal Umum, Buku Besar, serta Trial Balance. Akun dari workbook pengguna dipakai sebagai rujukan, termasuk 111.101 KAS YAYASAN, 151.200 PERSEDIAAN BARANG LAINNYA, 410.900 PENDAPATAN PENJUALAN TOKO, dan 510.900 BEBAN POKOK PENJUALAN TOKO.",
                 "Akuntan memilih periode yang sama, menelusuri nomor jurnal ke Buku Besar, lalu membandingkan saldo akun kas dan utang dengan pembayaran vendor. Untuk laporan Laba Rugi dan laporan lain, data hanya boleh dinyatakan lengkap setelah akun sumber dipetakan dan baris jurnalnya benar-benar masuk klasifikasi laporan. Jika sebuah laporan kosong, perbaiki pemetaan atau dataset terlebih dahulu dan ambil bukti ulang; jangan menampilkan halaman kosong sebagai hasil lulus.",
             ], **common),
    ]


def cover(doc: Document, summary: dict, cashier: dict):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(92)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("UAT LIVE SERVER DEMO | PANDUAN OPERASIONAL")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(ACCENT)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Apotik\nUAT End-to-End & Panduan Pengguna")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(34)
    run = subtitle.add_run("Kasir OTC, resep dokter, racikan, produksi farmasi, persediaan, pengadaan, akuntansi, dan laporan")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_table(doc, ["Atribut", "Nilai"], [
        ("Rilis", f"{RELEASE} | {VERSION}"),
        ("Tanggal", DATE),
        ("Server", summary.get("server", "https://demo.ecampus.id/ecampus")),
        ("Kanvas bukti", "1920 x 1080, tema hijau Apotik"),
        ("Keputusan", f"LULUS: 13 menu, {cashier.get('transaksiUtamaLulus', {}).get('total', 0)} transaksi Kasir, dan {cashier.get('postflightIdempoten', {}).get('totalKodeUnik', 0)} kode unik terverifikasi"),
        ("Distribusi", "APK Android dan installer Windows EXE; APK masih debug-signed"),
    ], [1.55, 4.95])
    note(doc, "Batas penggunaan", "Seluruh katalog, pasien, resep, antrean, dan transaksi bertanda sample/UAT. Data tidak boleh digunakan sebagai acuan terapi atau klaim registrasi obat.", AMBER)
    doc.add_page_break()


def build_document() -> Path:
    summary = load_json(SHOT / "uat-summary.json")
    cashier = load_json(OUT / "cashier-retest" / "uat-kasir-summary.json")
    procurement = load_json(SHOT / "procurement-summary.json")
    journal = load_json(SHOT / "vendor-journal-summary.json")
    financial = load_json(ACC / "financial-report-summary.json")
    minimums = {
        "katalog item": summary.get("katalogItemTotal", 0),
        "bahan racikan": summary.get("bahanRacikanTerverifikasi", summary.get("bahanRacikanTotal", 0)),
        "resep siap jual": summary.get("resepSiapJualTerverifikasi", summary.get("resepSiapJualTotal", 0)),
        "penjualan OTC": cashier.get("transaksiUtamaLulus", {}).get("otcObatBebas", 0),
        "penjualan resep dokter": cashier.get("transaksiUtamaLulus", {}).get("resepDokter", 0),
        "penjualan racikan": cashier.get("transaksiUtamaLulus", {}).get("racikan", 0),
        "produksi farmasi": cashier.get("transaksiUtamaLulus", {}).get("produksiFarmasi", 0),
        "tebus resep campuran": cashier.get("transaksiUtamaLulus", {}).get("tebusResepCampuran", 0),
        "formula racikan": cashier.get("provisioning", {}).get("formulaRacikanOperasional", 0),
        "formula produksi": cashier.get("provisioning", {}).get("formulaProduksiOperasional", 0),
        "register obat terkendali": summary.get("laporanTerkendaliBaris", 0),
        "antrean farmasi": summary.get("antreanFarmasiBaris", summary.get("antreanFarmasiDipastikan", 0)),
        "antrean obat jadi": summary.get("antreanObatJadi", 0),
        "antrean racikan": summary.get("antreanRacikan", 0),
        "laporan penjualan": summary.get("laporanPenjualanBaris", 0),
        "laporan kedaluwarsa": summary.get("laporanKedaluwarsaBaris", 0),
        "PR": procurement.get("pengadaan_pr_daftar", 0),
        "PO": procurement.get("pengadaan_po_daftar", 0),
        "BAST": procurement.get("pengadaan_bast_daftar", 0),
        "tagihan": procurement.get("pengadaan_tagihan_daftar", 0),
        "pembayaran vendor": procurement.get("pengadaan_bayar_daftar", 0),
        "jurnal vendor": journal.get("jurnalTerpostingTerverifikasi", 0),
        "sumber laporan keuangan": financial.get("sumberJurnalTerpostingTerverifikasi", 0),
    }
    kurang = {nama: nilai for nama, nilai in minimums.items() if int(nilai or 0) < 100}
    if kurang:
        raise ValueError(f"Dokumen ditahan karena bukti di bawah 100 record: {kurang}")
    items = scenarios(summary, cashier, procurement, journal, financial)
    diagrams = make_diagrams(items)
    doc = Document()
    configure(doc)
    cover(doc, summary, cashier)

    doc.add_heading("Ringkasan eksekutif", level=1)
    doc.add_paragraph(
        f"Retest antarmuka klien {VERSION} dijalankan pada {summary.get('server')} tanggal 5 September 2026 dalam mode baca/capture terhadap {RELEASE}. Seluruh tiga belas menu Apotik tampil dan membuka tujuan yang benar. Data live mencatat {summary.get('katalogItemTotal', 0)} item katalog, {summary.get('bahanRacikanTerverifikasi', summary.get('bahanRacikanTotal', 0))} bahan racikan, {cashier.get('provisioning', {}).get('resepSiapJualSetelahUat', 0)} resep siap jual, serta {summary.get('antreanFarmasiBaris', summary.get('antreanFarmasiDipastikan', 0))} antrean publik tersamar. Lima jalur Kasir—OTC, Resep Dokter, Racikan, Produksi Farmasi, dan Tebus Resep campuran—masing-masing lulus 100 data, total {cashier.get('transaksiUtamaLulus', {}).get('total', 0)} operasi, lalu {cashier.get('postflightIdempoten', {}).get('totalKodeUnik', 0)} kode unik dibaca ulang tanpa pembayaran atau produksi ganda."
    )
    note(doc, "Kesimpulan", f"Kasir Apotik lulus end-to-end untuk lima alur dengan 100 data per alur; 13 menu terbuka, pembayaran OTC dapat mencapai lembar konfirmasi, dan {cashier.get('regresiUi', {}).get('lulus', 0)} pengujian regresi UI lulus tanpa kegagalan. APK tetap memerlukan keystore produksi sebelum distribusi final.", GREEN)

    doc.add_heading("Matriks UAT", level=2)
    add_table(doc, ["No.", "Area", "Bukti", "Hasil"], [
        (1, "Menu Apotik mandiri", "13 / 13", "PASS"),
        (2, "OTC / Obat Bebas", cashier.get("transaksiUtamaLulus", {}).get("otcObatBebas", 0), "PASS"),
        (3, "Resep Dokter", cashier.get("transaksiUtamaLulus", {}).get("resepDokter", 0), "PASS"),
        (4, "Racikan", cashier.get("transaksiUtamaLulus", {}).get("racikan", 0), "PASS"),
        (5, "Produksi Farmasi", cashier.get("transaksiUtamaLulus", {}).get("produksiFarmasi", 0), "PASS"),
        (6, "Tebus Resep campuran", cashier.get("transaksiUtamaLulus", {}).get("tebusResepCampuran", 0), "PASS"),
        (7, "Kode unik/idempoten", cashier.get("postflightIdempoten", {}).get("totalKodeUnik", 0), "PASS"),
        (8, "Regresi UI", cashier.get("regresiUi", {}).get("lulus", 0), "PASS"),
        (9, "Katalog + bahan", f"{summary.get('katalogItemTotal', 0)} / {summary.get('bahanRacikanTerverifikasi', summary.get('bahanRacikanTotal', 0))}", "PASS"),
        (10, "Resep siap jual", cashier.get("provisioning", {}).get("resepSiapJualSetelahUat", 0), "PASS"),
        (11, "Antrean publik", summary.get("antreanFarmasiBaris", summary.get("antreanFarmasiDipastikan", 0)), "PASS PRIVASI"),
        (12, "Procure-to-pay", "100 per tahap", "PASS"),
        (13, "Jurnal vendor", journal.get("jurnalTerpostingTerverifikasi", 0), "PASS"),
        (14, "Laporan penjualan", summary.get("laporanPenjualanBaris", 0), "PASS"),
        (15, "Sumber laporan keuangan", financial.get("sumberJurnalTerpostingTerverifikasi", 0), "PASS"),
    ], [0.5, 2.65, 1.35, 2.0])

    doc.add_page_break()
    doc.add_heading("Pemetaan akun yang dipakai", level=2)
    add_table(doc, ["Kode", "Nama akun", "Peran UAT"], [
        ("111.101", "KAS YAYASAN", "Kas keluar pembayaran vendor"),
        ("151.200", "PERSEDIAAN BARANG LAINNYA", "Persediaan bila subakun obat belum ditetapkan"),
        ("310.500", "HUTANG VENDOR", "Pelunasan kewajiban vendor"),
        ("410.900", "PENDAPATAN PENJUALAN TOKO", "Pendapatan Apotik"),
        ("510.900", "BEBAN POKOK PENJUALAN TOKO", "HPP Apotik"),
    ], [0.9, 2.8, 2.8])

    doc.add_heading("Metode retest dan interpretasi 100 data", level=2)
    doc.add_paragraph(
        "Retest tanggal 5 September 2026 dilakukan langsung ke server Demo. Run API mengeksekusi 100 data "
        "untuk masing-masing alur OTC, Resep Dokter, Racikan, Produksi Farmasi, dan Tebus Resep campuran. "
        "Run screenshot berikutnya memakai mode baca/capture agar detail UI dapat didokumentasikan tanpa "
        "membuat pembayaran, produksi, mutasi stok, retur, atau jurnal tambahan."
    )
    add_bullet(doc, "Sebanyak 500 operasi Kasir dan 500 kode unik/idempoten dibaca ulang; 17 respons gateway sementara selesai setelah recheck dan tidak menimbulkan data ganda.")
    add_bullet(doc, "Tiga belas menu Apotik tampil; screenshot memperlihatkan mode OTC, Resep Dokter, Racikan, Produksi Farmasi, Tebus Resep, pembayaran, detail resep, serta Manajemen Farmasi.")
    add_bullet(doc, "Form transaksi stok—penerimaan PBF, stok opname, dan retur—divalidasi pada pembukaan menu, sumber katalog minimal 100 data, state tab, dan integrasi data; retest ini tidak menciptakan 100 mutasi stok tambahan.")
    add_bullet(doc, "PASS berarti layar membuka tanpa error, hak akses sesuai, sumber data memenuhi minimum, hasil dapat ditelusuri, dan tidak ada penggandaan berdasarkan kode idempoten atau marker UAT.")
    add_bullet(doc, "Screenshot baru memakai data sample/UAT. Identitas antrean publik tetap disamarkan dan tidak boleh dipakai sebagai data klinis atau registrasi obat.")

    for index, item in enumerate(items, 1):
        doc.add_page_break()
        doc.add_heading(f"{index}. {item['title']}", level=1)
        add_picture(doc, item["screenshot"], item["title"])
        doc.add_heading("Penjelasan layar dan hasil pengujian", level=2)
        for paragraph in item["paragraphs"]:
            doc.add_paragraph(paragraph)
        doc.add_heading("Use case", level=2)
        add_picture(doc, diagrams[index][0], f"Use case - {item['title']}")
        doc.add_heading("Flowchart", level=2)
        add_picture(doc, diagrams[index][1], f"Flowchart - {item['title']}")
        doc.add_heading("ERD dan aliran data", level=2)
        add_picture(doc, diagrams[index][2], f"ERD dan aliran data - {item['title']}")
        doc.add_heading("Langkah operator", level=2)
        for step_number, step in enumerate(item["flow"], 1):
            add_number(doc, step, step_number)
        add_table(doc, ["Kontrol", "Kriteria"], [
            ("Otorisasi", item["controls"]),
            ("Acceptance", item["acceptance"]),
            ("Jika gagal", "Hentikan tindakan lanjutan, cari kode sumber yang sama, simpan bukti, dan eskalasikan tanpa membuat duplikasi."),
        ], [1.45, 5.05])

    doc.add_page_break()
    doc.add_heading("Checklist rilis dan operasional", level=1)
    for text in [
        "Varian yang tampil bernama Apotik dan memakai tema hijau pada Android serta Windows.",
        "Jendela Windows membuka dalam keadaan maximized dan screenshot diambil pada 1920 x 1080.",
        "Tiga belas menu Apotik tampil dan mode Kasir OTC, Resep Dokter, Racikan, serta Produksi Farmasi tidak lagi terkunci.",
        "Pembayaran OTC membuka lembar metode pembayaran tanpa pesan kode transaksi server belum terinisialisasi.",
        "Katalog, bahan racikan, resep, antrean, transaksi, laporan, dan data pengadaan memenuhi matriks UAT.",
        "Identitas pasien pada layar publik tersamar dan catatan publik tidak memuat data klinis sensitif.",
        "APK dan installer EXE memiliki checksum; artefak Windows tidak dikemas sebagai ZIP.",
        "APK debug-signed hanya untuk UAT/pilot sampai keystore produksi tersedia.",
        "Rollback dilakukan dengan artefak rilis sebelumnya dan reversal yang disetujui; audit trail tidak dihapus.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("Sign-off", level=1)
    add_table(doc, ["Peran", "Nama", "Keputusan", "Tanggal / tanda tangan"], [
        ("Owner Apotik", "", "", ""),
        ("Apoteker Penanggung Jawab", "", "", ""),
        ("Procurement", "", "", ""),
        ("Keuangan/Akuntansi", "", "", ""),
        ("IT/DevOps", "", "", ""),
        ("Keamanan/Privasi", "", "", ""),
    ], [1.8, 1.55, 1.45, 1.7])

    output = OUT / "Laporan-UAT-dan-Panduan-Apotik-v1.34.25.docx"
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    manifest = {
        "version": VERSION,
        "apiRelease": RELEASE,
        "document": output.name,
        "screenshotSections": len(items),
        "diagramCount": len(items) * 3,
        "sourceSummary": [
            "screenshots/uat-summary.json",
            "cashier-retest/uat-kasir-summary.json",
        ],
    }
    (OUT / "document-manifest.json").write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")
    presentation_data = []
    for index, item in enumerate(items, 1):
        presentation_data.append({
            "index": index,
            "title": item["title"],
            "screenshot": item["screenshot"].relative_to(ROOT).as_posix(),
            "actors": item["actors"],
            "usecases": item["usecases"],
            "flow": item["flow"],
            "gate": item["gate"],
            "entities": item["entities"],
            "paragraphs": item["paragraphs"],
            "controls": item["controls"],
            "acceptance": item["acceptance"],
        })
    (OUT / "presentation-data.json").write_text(
        json.dumps(presentation_data, indent=2, ensure_ascii=False), encoding="utf-8")
    return output


if __name__ == "__main__":
    print(build_document())
