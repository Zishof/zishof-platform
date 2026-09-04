from __future__ import annotations

import json
import re
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "pos-apotik-emedik" / "uat-v1.34.23"
SHOT = OUT / "screenshots"
PROC = OUT / "screenshots-pengadaan"
ACC = OUT / "screenshots-akuntansi"
DIAGRAM = OUT / "diagrams"
VERSION = "1.34.23 (build 185)"
RELEASE = "apotik-v1.34.23"
DATE = "4 September 2026"

NAVY = "0B1F33"
BLUE = "075985"
TEAL = "0F766E"
GREEN = "15803D"
AMBER = "B45309"
RED = "B91C1C"
INK = "172033"
LIGHT = "F4F7FA"
MID = "D9E5EC"


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def _center_text(draw: ImageDraw.ImageDraw, xy, text: str, size=23, bold=True,
                 color="#172033", max_chars=24):
    lines: list[str] = []
    for part in text.split("\n"):
        lines.extend(wrap(part, max_chars) or [""])
    fnt = font(size, bold)
    heights = []
    for line in lines:
        box = draw.textbbox((0, 0), line, font=fnt)
        heights.append(box[3] - box[1] + 7)
    y = xy[1] + max(12, ((xy[3] - xy[1]) - sum(heights)) / 2)
    for line, height in zip(lines, heights):
        box = draw.textbbox((0, 0), line, font=fnt)
        x = (xy[0] + xy[2] - (box[2] - box[0])) / 2
        draw.text((x, y), line, font=fnt, fill=color)
        y += height


def _box(draw, xy, label, fill="#FFFFFF", outline="#075985", size=22):
    draw.rounded_rectangle(xy, radius=20, fill=fill, outline=outline, width=3)
    _center_text(draw, xy, label, size=size, max_chars=25)


def _arrow(draw, start, end, color="#075985", width=5):
    draw.line([start, end], fill=color, width=width)
    x, y = end
    if abs(end[0] - start[0]) >= abs(end[1] - start[1]):
        direction = 1 if end[0] > start[0] else -1
        draw.polygon([(x, y), (x - direction * 18, y - 10),
                      (x - direction * 18, y + 10)], fill=color)
    else:
        direction = 1 if end[1] > start[1] else -1
        draw.polygon([(x, y), (x - 10, y - direction * 18),
                      (x + 10, y - direction * 18)], fill=color)


def _canvas(title: str):
    image = Image.new("RGB", (1600, 900), "#F7FAFC")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1600, 96), fill=f"#{NAVY}")
    draw.text((48, 26), title, font=font(33, True), fill="white")
    return image, draw


def make_use_case(index: int, item: dict) -> Path:
    image, draw = _canvas(f"Use Case {index} — {item['title']}")
    actors = item["actors"]
    actor_positions = [(45, 160, 280, 265), (45, 370, 280, 475),
                       (45, 580, 280, 685), (1320, 365, 1555, 480)]
    actor_colors = ["#DBEAFE", "#EDE9FE", "#FEF3C7", "#DCFCE7"]
    for pos, actor, color in zip(actor_positions, actors, actor_colors):
        _box(draw, pos, actor, color, "#075985", 21)
    cases = item["usecases"]
    case_positions = [(430, 145, 1165, 250), (430, 310, 1165, 415),
                      (430, 475, 1165, 580), (430, 640, 1165, 745)]
    for pos, label in zip(case_positions, cases):
        _box(draw, pos, label, "#FFFFFF", "#0F766E", 22)
    for idx, pos in enumerate(actor_positions[:3]):
        target = case_positions[min(idx, 2)]
        _arrow(draw, (pos[2], (pos[1] + pos[3]) // 2),
               (target[0], (target[1] + target[3]) // 2))
    _arrow(draw, (actor_positions[3][0], 422), (case_positions[3][2], 692), "#15803D")
    draw.text((48, 825), "Garis menunjukkan interaksi utama; otorisasi aktual tetap mengikuti RBAC dan SOP institusi.",
              font=font(20), fill="#475569")
    path = DIAGRAM / f"{index:02d}-use-case.png"
    image.save(path)
    return path


def make_flow(index: int, item: dict) -> Path:
    image, draw = _canvas(f"Flowchart {index} — {item['title']}")
    positions = [(50, 190, 350, 315), (455, 190, 755, 315),
                 (860, 190, 1160, 315), (1265, 190, 1550, 315),
                 (860, 540, 1160, 665), (455, 540, 755, 665)]
    fills = ["#DBEAFE", "#FFFFFF", "#FFFFFF", "#FEF3C7", "#DCFCE7", "#EDE9FE"]
    outlines = ["#2563EB", "#075985", "#0F766E", "#D97706", "#15803D", "#7C3AED"]
    for pos, label, fill, outline in zip(positions, item["flow"], fills, outlines):
        _box(draw, pos, label, fill, outline, 21)
    for left, right in zip(positions[:3], positions[1:4]):
        _arrow(draw, (left[2], 252), (right[0], 252))
    _arrow(draw, (1405, positions[3][3]), (1010, positions[4][1]), "#B45309")
    _arrow(draw, (positions[4][0], 602), (positions[5][2], 602), "#15803D")
    draw.rounded_rectangle((310, 760, 1290, 835), radius=18, fill="#FFF7ED", outline="#D97706", width=2)
    _center_text(draw, (310, 760, 1290, 835), item["gate"], size=19, max_chars=92)
    path = DIAGRAM / f"{index:02d}-flowchart.png"
    image.save(path)
    return path


def make_erd(index: int, item: dict) -> Path:
    image, draw = _canvas(f"ERD dan Aliran Data {index} — {item['title']}")
    positions = [(70, 170, 390, 315), (640, 150, 960, 295),
                 (1210, 170, 1530, 315), (330, 530, 650, 675),
                 (950, 530, 1270, 675)]
    fills = ["#DBEAFE", "#FFFFFF", "#FEF3C7", "#EDE9FE", "#DCFCE7"]
    outlines = ["#2563EB", "#0F766E", "#D97706", "#7C3AED", "#15803D"]
    for pos, entity, fill, outline in zip(positions, item["entities"], fills, outlines):
        _box(draw, pos, entity, fill, outline, 20)
    links = [((390, 242), (640, 222)), ((960, 222), (1210, 242)),
             ((800, 295), (490, 530)), ((800, 295), (1110, 530)),
             ((650, 602), (950, 602))]
    for start, end in links:
        _arrow(draw, start, end, "#075985", 4)
    labels = [(480, 175, "1..n"), (1070, 175, "1..n"),
              (540, 405, "jejak audit"), (960, 405, "posting"),
              (755, 560, "rekonsiliasi")]
    for x, y, label in labels:
        draw.text((x, y), label, font=font(18, True), fill="#475569")
    draw.text((48, 820), "Model konseptual UAT; nama tabel fisik dapat berbeda. Kunci referensi wajib menjaga dokumen sumber dan audit trail.",
              font=font(20), fill="#475569")
    path = DIAGRAM / f"{index:02d}-erd-dataflow.png"
    image.save(path)
    return path


def make_diagrams(items: list[dict]) -> dict[int, tuple[Path, Path, Path]]:
    DIAGRAM.mkdir(parents=True, exist_ok=True)
    return {idx: (make_use_case(idx, item), make_flow(idx, item), make_erd(idx, item))
            for idx, item in enumerate(items, 1)}


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, value, bold=False, color=INK, size=9):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(value))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def add_table(doc: Document, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        set_cell(table.rows[0].cells[idx], header, True, "FFFFFF", 9)
        shade(table.rows[0].cells[idx], BLUE)
        if widths:
            table.rows[0].cells[idx].width = Inches(widths[idx])
    set_repeat_header(table.rows[0])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value, False, INK, 8.5)
            if row_idx % 2:
                shade(cells[idx], "F8FAFC")
            if widths:
                cells[idx].width = Inches(widths[idx])
    return table


def add_page_field(paragraph):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Title", 28, NAVY, 0, 14),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = f"POS APOTIK • UAT & MANUAL • {VERSION}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Calibri"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string("64748B")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Dokumen terkendali • Halaman ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    add_page_field(footer)

    doc.core_properties.title = "Laporan UAT dan Manual POS Apotik v1.34.23"
    doc.core_properties.subject = "UAT end-to-end server demo, data volume, akuntansi, laporan, dan panduan operasi"
    doc.core_properties.author = "Tim eBisnis / Zishof"
    doc.core_properties.keywords = "POS Apotik, UAT, manual, farmasi, pengadaan, akuntansi, Android, Windows"


def cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(65)
    r = p.add_run("DOKUMEN PENERIMAAN PENGGUNA • DATA REAL SERVER DEMO")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(TEAL)

    p = doc.add_paragraph(style="Title")
    p.add_run("UAT End-to-End & Manual Operasional\nPOS Apotik")
    p.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    r = p.add_run("Penjualan obat, racikan, pengadaan, stok, kedaluwarsa, jurnal, posting, dan laporan keuangan")
    r.font.name = "Calibri"
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string("475569")
    p.paragraph_format.space_after = Pt(24)

    add_table(doc, ["Atribut", "Nilai"], [
        ("Rilis", f"{RELEASE} • {VERSION}"),
        ("Tanggal uji", DATE),
        ("Target", "https://demo.ecampus.id/ecampus/"),
        ("Volume", "100 transaksi per skenario; 100 dokumen per tahap pengadaan"),
        ("Keputusan", "LULUS BERSYARAT — layak UAT/pilot terkontrol, belum produksi tersign"),
        ("Artefak", "Word, PDF, PPTX, APK Android, installer Windows EXE"),
    ], [1.55, 4.85])
    doc.add_paragraph()
    p = doc.add_paragraph("Kerahasiaan dan keselamatan: data pasien pada layar publik harus tersamar. Data uji tidak boleh digunakan sebagai dasar terapi klinis. Dokumen ini membedakan PASS fungsional, PASS tampilan, dan blocker integrasi agar sign-off tidak menyesatkan.")
    p.style = doc.styles["Caption"]
    doc.add_page_break()


def note(doc: Document, label: str, text: str, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, "EAF3F8")
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    paragraph.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.187)
    paragraph.add_run(text)


def add_number(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.187)
    paragraph.add_run(text)


def add_picture(doc: Document, path: Path, caption: str, width=6.45):
    if not path.exists():
        note(doc, "Bukti tidak ditemukan", str(path), RED)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    paragraph = doc.add_paragraph(caption)
    paragraph.style = doc.styles["Caption"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def word_count(paragraphs: list[str]) -> int:
    return len(re.findall(r"\b[\wÀ-ÿ-]+\b", " ".join(paragraphs), flags=re.UNICODE))


def narrative(item: dict) -> list[str]:
    title = item["title"]
    actors = ", ".join(item["actors"])
    steps = "; ".join(item["flow"])
    entities = ", ".join(item["entities"])
    paragraphs = [
        f"Bukti layar untuk {title} dipakai sebagai titik masuk pemeriksaan UAT sekaligus panduan kerja operator. Tujuan pengujian bukan hanya memastikan halaman dapat terbuka, melainkan memastikan rangkaian keputusan, data, otorisasi, dan jejak audit tetap konsisten dari awal sampai akhir. Aktor yang terlibat adalah {actors}. Sebelum memulai, penguji memastikan varian yang tampil adalah POS Apotik, koneksi HTTPS menuju server demo aktif, pengguna berhasil login dengan role yang sesuai, tanggal dan waktu perangkat benar, serta data uji mempunyai penanda rilis agar dapat dibedakan dari transaksi lain pada tenant bersama. Screenshot diambil pada resolusi 1600 × 900 setelah indikator pemuatan selesai. Karena server demo dipakai bersama, jumlah total yang tampak pada dashboard dapat melebihi volume yang dibuat dalam run ini; hasil kelulusan karena itu dihitung dari kode idempoten dan ringkasan API, bukan dari perkiraan visual semata.",
        f"Cara penggunaan normal mengikuti urutan: {steps}. Operator membaca seluruh label dan nilai sebelum menekan tindakan simpan atau setujui. Kolom wajib tidak boleh diisi dengan nilai acak hanya untuk melewati validasi. Untuk nominal, kuantitas, batch, tanggal, vendor, pasien, dan akun, penguji mencocokkan kembali sumbernya dengan dokumen atau skenario. Setelah penyimpanan, operator menunggu notifikasi sukses, kemudian membuka ulang daftar dan mencari kode uji yang sama. Langkah baca ulang ini penting untuk membuktikan persistensi, bukan sekadar perubahan state sementara pada layar. Jika tombol tidak aktif, aplikasi menampilkan pesan gagal, atau indikator terus berputar, penguji mencatat waktu, menu, kode dokumen, payload relevan tanpa kredensial, dan screenshot. Penguji tidak mengulang simpan dengan kode baru sebelum memastikan apakah permintaan pertama sebenarnya sudah diterima server, karena tindakan tersebut dapat menggandakan transaksi.",
        f"Data yang diamati dalam skenario ini membentuk aliran konseptual antara {entities}. Hubungan antardata dijaga melalui identitas dokumen, referensi transaksi sumber, detail barang atau akun, dan status proses. Kunci idempoten pada run volume mencegah penciptaan ulang ketika tes diulang setelah gangguan jaringan. Untuk data persediaan, kuantitas tidak boleh hanya dicatat di header; detail batch, tanggal kedaluwarsa, dan sisa stok harus dapat ditelusuri. Untuk data keuangan, jurnal harus membawa referensi dokumen sumber dan pasangan debit-kredit yang seimbang. Untuk data pasien, layar petugas boleh menampilkan informasi operasional yang diperlukan, sedangkan layar publik wajib menggunakan identitas tersamar. Penguji juga memperhatikan bahwa keberhasilan respons HTTP atau status “success” belum selalu membuktikan angka laporan benar; rekonsiliasi nilai terhadap dokumen sumber tetap menjadi acceptance test terpisah.",
        f"Hasil aktual pada server demo untuk skenario ini adalah: {item['observation']}. Status formalnya {item['status']}. Arti status tersebut harus dibaca secara presisi. PASS berarti operasi yang dinyatakan berhasil benar-benar disimpan atau dirender dan diverifikasi dengan pemeriksaan lanjutan. PASS TAMPILAN hanya berarti komponen dapat digunakan dan difoto, bukan bahwa seluruh endpoint produksi sudah tersedia. BLOCKED berarti prasyarat backend, konfigurasi, atau data master belum mencukupi; blocker tidak boleh disamarkan sebagai kegagalan operator dan juga tidak boleh dianggap lulus. LULUS BERSYARAT berarti jalur utama dapat dipakai untuk UAT atau pilot yang dikendalikan, dengan risiko dan pekerjaan retest yang tercantum. Pemisahan ini menjaga dokumen tetap dapat diaudit dan membantu pemilik proses memutuskan apakah sebuah temuan menahan produksi atau cukup ditangani sebelum go-live.",
        f"Kontrol utama untuk {title} adalah {item['controls']}. Penguji melakukan negative check dengan mempertimbangkan data kosong, hak akses tidak memadai, dokumen sudah diproses, batch tidak layak, ketidakseimbangan jurnal, duplikasi kode, dan putus jaringan. Sistem harus menolak atau membatasi tindakan yang tidak sah tanpa menghapus data yang telah benar. Setelah setiap perubahan status, petugas yang berbeda—bila segregasi tugas diwajibkan—membuka data dari sesinya sendiri dan mengonfirmasi hasil. Bukti ideal terdiri dari screenshot sebelum tindakan, respons sesudah tindakan, daftar hasil, dan ringkasan mesin. Dokumen ini menampilkan satu screenshot representatif per kelompok skenario agar setiap gambar dapat diberi narasi dan diagram lengkap; berkas screenshot lain tetap tersedia di folder bukti dalam repository untuk investigasi teknis, namun tidak ditampilkan di badan dokumen agar tidak menimbulkan interpretasi bahwa semua gambar memiliki status integrasi yang sama.",
        f"Bila terjadi penyimpangan, operator menjalankan prosedur berikut. Pertama, hentikan tindakan lanjutan yang dapat mengubah stok atau keuangan. Kedua, catat kode dokumen, waktu, role, dan pesan kesalahan. Ketiga, muat ulang daftar serta cari kode idempoten yang sama untuk menentukan apakah server sudah menyimpan transaksi. Keempat, bandingkan detail dengan dokumen sumber dan jangan membuat pengganti sebelum status transaksi pertama pasti. Kelima, eskalasikan ke pemilik modul dengan klasifikasi: data master, RBAC, validasi bisnis, integrasi API, tampilan, atau akuntansi. Jika penyimpangan menyentuh pasien, obat, batch, atau kedaluwarsa, apoteker mengambil keputusan keselamatan dan menahan penyerahan bila perlu. Jika menyentuh vendor, tagihan, pembayaran, atau jurnal, keuangan menahan pembayaran/posting dan akuntansi melakukan rekonsiliasi. Seluruh koreksi harus menghasilkan jejak audit baru, bukan menghapus diam-diam riwayat lama.",
        f"Kriteria penerimaan akhir untuk skenario ini adalah: {item['acceptance']}. Setelah kriteria terpenuhi, penguji memberi tanda pada matriks, mencatat bukti, dan meminta sign-off pemilik proses. Untuk penggunaan harian, supervisor sebaiknya menjalankan pemeriksaan awal shift terhadap koneksi, toko aktif, perangkat periferal, role, dan sinkronisasi; pemeriksaan selama shift terhadap transaksi tertahan, stok negatif, batch mendekati kedaluwarsa, antrean, dan dokumen gagal; serta pemeriksaan akhir shift terhadap kas, stok kritis, dokumen pengadaan, pembayaran, jurnal, dan laporan. Pola ini membuat aplikasi berfungsi sebagai sistem pengendalian, bukan hanya alat input. Pada rilis {VERSION}, kesimpulan {item['status']} berlaku untuk lingkungan yang diuji dan tidak otomatis membuktikan kompatibilitas setiap tenant, printer, monitor, kebijakan COA, atau keystore. Retest produksi wajib memakai konfigurasi tujuan dan persetujuan owner.",
    ]
    assert word_count(paragraphs) >= 500, (title, word_count(paragraphs))
    return paragraphs


def items() -> list[dict]:
    return [
        dict(title="Penjualan Obat Jadi", screenshot=SHOT / "01-kasir-obat-jadi.png",
             caption="Bukti 1. Kasir obat jadi setelah UAT API pada server demo.",
             actors=["Pasien / keluarga", "Kasir", "Apoteker", "Supervisor"],
             usecases=["Identifikasi pembeli dan permintaan", "Cari obat serta validasi hak jual", "Pilih batch FEFO dan proses pembayaran", "Edukasi, serah terima, dan audit"],
             flow=["Identifikasi pasien", "Cari/pindai obat", "Validasi stok & golongan", "Pilih batch FEFO", "Bayar & simpan", "Etiket dan serah terima"],
             gate="GATE: obat belum kedaluwarsa, jumlah cukup, otorisasi sesuai, kode transaksi idempoten.",
             entities=["PASIEN\nid, identitas", "PENJUALAN\nkode, tanggal", "DETAIL OBAT\nqty, harga", "BATCH\nED, sisa", "JURNAL/AUDIT\nreferensi"],
             observation="100 dari 100 transaksi obat jadi dengan kode UAT v1.34.23 berhasil dan pemeriksaan idempotensi lulus. Katalog deploy hanya memuat dua item, sehingga keberhasilan transaksi volume tidak berarti katalog 100 item telah tersedia",
             status="PASS transaksi / BLOCKED volume katalog", controls="validasi item, batch tidak kedaluwarsa, sisa stok, FEFO, obat terkendali, pembayaran, dan bukti serah terima",
             acceptance="100 kode tersimpan satu kali, stok berkurang pada batch yang benar, transaksi dapat dibaca ulang, dan penyerahan dilakukan setelah verifikasi pasien-obat-jumlah"),
        dict(title="Penjualan Obat Racik", screenshot=SHOT / "02-resep-racikan.png",
             caption="Bukti 2. Halaman resep/racikan dengan data server demo.",
             actors=["Dokter / resep", "Apoteker", "Tenaga teknis", "Pasien"],
             usecases=["Terima dan telaah resep", "Hitung serta pilih bahan", "Racik, kemas, dan periksa akhir", "Edukasi serta ubah status antrean"],
             flow=["Resep diterima", "Telaah klinis/farmasetik", "Hitung bahan", "Racik & kemas", "Pemeriksaan akhir", "Serahkan & edukasi"],
             gate="GATE: resep terbaca, bahan tersedia, kalkulasi dan etiket diverifikasi petugas berwenang.",
             entities=["RESEP\nid, pasien", "RACIKAN\nformula, dosis", "BAHAN\nitem, qty", "BATCH BAHAN\nED, sisa", "PENYERAHAN\nstatus, edukasi"],
             observation="100 dari 100 transaksi racikan berhasil; endpoint daftar mengembalikan 100 resep pada halaman uji dan ringkasan provision server menyatakan total 5.000 resep demo tersedia",
             status="PASS transaksi racikan", controls="telaah resep, kalkulasi dosis, kompatibilitas, pemilihan batch, kebersihan area racik, pemeriksaan akhir, etiket, dan edukasi",
             acceptance="100 resep terhubung ke transaksi satu kali, pemakaian bahan dapat ditelusuri, status selesai benar, dan hasil racikan diserahkan setelah pemeriksaan akhir"),
        dict(title="Penjualan Gabungan dan Layar Pasien", screenshot=SHOT / "06-layar-kedua-obat-jadi-racikan.png",
             caption="Bukti 3. Pratinjau terkontrol layar kedua gabungan; endpoint antrean live belum tersedia.",
             actors=["Kasir", "Apoteker", "Pasien / keluarga", "Layar publik"],
             usecases=["Gabungkan obat jadi dan racikan", "Bayar dalam satu transaksi", "Perbarui status penyiapan", "Tampilkan antrean tersamar"],
             flow=["Pilih item jadi", "Pilih resep racikan", "Validasi dua lini", "Bayar gabungan", "Siapkan obat", "Panggil pada layar"],
             gate="GATE: seluruh lini sah; layar publik hanya menerima nomor antrean, status, loket, dan identitas tersamar.",
             entities=["PENJUALAN\nkode gabungan", "DETAIL JADI\nitem, batch", "DETAIL RACIK\nresep, bahan", "ANTREAN\nnomor, status", "LAYAR PUBLIK\nview tersamar"],
             observation="100 dari 100 transaksi gabungan yang masing-masing berisi lini obat jadi dan racikan berhasil. Komponen layar gabungan, khusus obat jadi, dan khusus racikan dapat dirender, tetapi endpoint antrean farmasi pada server demo mengembalikan aksi belum tersedia; screenshot ini memakai data pratinjau terkontrol",
             status="PASS transaksi / PASS komponen / BLOCKED integrasi antrean", controls="atomicity transaksi gabungan, pencegahan duplikasi, sinkronisasi status, privasi pasien, polling, dan pemilihan monitor",
             acceptance="100 transaksi gabungan tersimpan utuh, tidak ada lini terputus, dan integrasi layar baru lulus setelah endpoint live menampilkan minimal 50 antrean tersamar serta perubahan status real-time"),
        dict(title="Kulakan: PR, PO, BAST, Tagihan, dan Pembayaran", screenshot=PROC / "02-pr-daftar-50.png",
             caption="Bukti 4. Dashboard PR pada rangkaian procure-to-pay server demo.",
             actors=["Peminta", "Approver", "Gudang", "Keuangan"],
             usecases=["Ajukan kebutuhan melalui PR", "Terbitkan PO dan terima barang", "Cocokkan BAST dengan tagihan", "Bayar vendor dan buat audit trail"],
             flow=["PR + justifikasi", "Persetujuan", "PO ke vendor", "BAST & batch", "Terima tagihan", "Bayar vendor"],
             gate="GATE: otorisasi dan three-way match PR/PO/BAST/tagihan konsisten sebelum pembayaran.",
             entities=["PR\nkode, kebutuhan", "PO\nvendor, termin", "BAST\nqty, mutu", "TAGIHAN\nnomor, nilai", "PEMBAYARAN\nbukti, jurnal"],
             observation="100 PR, 100 PO, 100 BAST, dan 100 tagihan dengan prefix rilis berhasil dibaca. Seratus pembayaran berhasil disetujui dan dijadikan sumber jurnal, namun filter daftar pembayaran memakai marker yang sama mengembalikan nol sehingga query/filter tersebut dicatat sebagai defect. Dashboard bersama menampilkan total lebih besar dari run ini",
             status="PASS alur P2P / defect filter daftar pembayaran", controls="segregasi tugas, otorisasi anggaran, vendor, qty, mutu, batch, ED, three-way match, bukti transfer, dan pelacakan status",
             acceptance="100 dokumen per tahap terhubung konsisten, tidak ada pembayaran tanpa tagihan disetujui, serta filter pembayaran diperbaiki dan diretest terhadap marker rilis"),
        dict(title="Pengelolaan Stok Semua Obat dan Alat Kesehatan", screenshot=SHOT / "03-formularium-obat.png",
             caption="Bukti 5. Formularium/katalog Apotik yang tersedia pada deployment demo.",
             actors=["Petugas gudang", "Apoteker", "Kasir", "Auditor stok"],
             usecases=["Kelola master dan kategori", "Terima batch serta mutasi", "Reservasi/pengeluaran FEFO", "Stock opname dan rekonsiliasi"],
             flow=["Buat/validasi item", "Kelompokkan kategori", "Terima batch", "Mutasi/penjualan", "Hitung fisik", "Rekonsiliasi selisih"],
             gate="GATE: satuan, kategori, batch, lokasi, sisa, dan tanggal kedaluwarsa tervalidasi sebelum item aktif.",
             entities=["ITEM\nkode, kategori", "BATCH\nlot, ED", "MUTASI\nmasuk/keluar", "LOKASI\nrak, gudang", "OPNAME\nfisik, selisih"],
             observation="pencarian katalog pada deployment demo hanya mengembalikan dua item dan satu item non-terkendali yang layak dipakai untuk uji transaksi. Provision telah dijalankan, tetapi volume master obat yang diminta belum muncul di endpoint katalog. Uji 100 transaksi tetap dapat dilakukan melalui top-up batch idempoten",
             status="BLOCKED data master volume / PASS fungsi stok yang terpakai", controls="keunikan kode, kategori obat jadi, bahan racikan, alat kesehatan, satuan, batch, FEFO, stok negatif, mutasi, dan stock opname",
             acceptance="minimal 100 item lintas kategori tersedia dan dapat ditelusuri per batch/lokasi; mutasi dari penerimaan, penjualan, racikan, retur, dan opname menghasilkan saldo yang dapat direkonsiliasi"),
        dict(title="Pengelolaan Tanggal Kedaluwarsa", screenshot=SHOT / "04-batch-kedaluwarsa.png",
             caption="Bukti 6. Monitor batch dan kedaluwarsa pada POS Apotik.",
             actors=["Sistem", "Gudang", "Apoteker", "Supervisor"],
             usecases=["Catat lot dan tanggal ED", "Prioritaskan batch FEFO", "Beri peringatan mendekati ED", "Karantina/retur/musnahkan"],
             flow=["Terima lot", "Validasi ED", "Simpan batch", "Pantau horizon", "Karantina jika tidak layak", "Retur/musnahkan"],
             gate="GATE: batch kedaluwarsa tidak boleh dipilih untuk penjualan atau bahan racikan.",
             entities=["ITEM\nkode", "BATCH\nlot, ED", "ALERT\nhorizon, status", "KARANTINA\nalasan", "DISPOSISI\nretur/musnah"],
             observation="monitor batch dan laporan kedaluwarsa berhasil diakses; run akhir mendeteksi setidaknya satu batch kedaluwarsa ketika memilih batch uji dan laporan API mengembalikan baris data. Algoritme kemudian memilih batch aktif dengan sisa tertinggi untuk mencegah false failure akibat batch FEFO yang hampir habis",
             status="PASS monitoring / perlu sign-off SOP disposisi", controls="format tanggal, zona waktu, horizon peringatan, FEFO, larangan batch lewat ED, karantina, otorisasi retur/pemusnahan, dan audit trail",
             acceptance="batch kedaluwarsa selalu ditolak dari transaksi, batch mendekati ED tampil dalam monitoring, dan setiap disposisi mempunyai alasan, jumlah, petugas, persetujuan, serta bukti"),
        dict(title="Akuntansi: Jurnal dan Posting", screenshot=ACC / "08-draft-jurnal.png",
             caption="Bukti 7. Dashboard Draft Jurnal setelah pengujian vendor dan posting.",
             actors=["Keuangan", "Akuntan", "Approver", "Auditor"],
             usecases=["Bangkitkan/masukkan jurnal", "Validasi debit dan kredit", "Posting ke buku besar", "Rekonsiliasi dengan sumber"],
             flow=["Dokumen sumber", "Preview/mapping akun", "Jurnal draf", "Validasi seimbang", "Posting", "Rekonsiliasi"],
             gate="GATE: debit = kredit, periode terbuka, akun sah, referensi sumber dan approval tersedia.",
             entities=["DOKUMEN SUMBER\npembayaran", "JURNAL\ntanggal, referensi", "DETAIL JURNAL\nakun, D/K", "BUKU BESAR\nposting", "AUDIT\nuser, waktu"],
             observation="100 pembayaran vendor diperiksa, 100 jurnal dibuat, 100 diposting pada run, dan 100 diverifikasi terposting. Mapping UAT memakai debit 310.500 HUTANG VENDOR dan kredit 111.101 KAS YAYASAN. Dashboard tenant bersama kemudian menampilkan 777 draf, 151 posted, total 928; angka dashboard tidak semuanya berasal dari run ini",
             status="PASS jurnal vendor / BLOCKED posting otomatis tertentu", controls="pemetaan COA, keseimbangan, periode, approval, referensi sumber, pencegahan posting ganda, reversal, dan rekonsiliasi",
             acceptance="100 jurnal berstatus posted satu kali, total debit sama dengan kredit, akun sesuai workbook, dan buku besar dapat menelusuri kembali pembayaran vendor"),
        dict(title="Laporan Penjualan dan Kulakan", screenshot=SHOT / "09-laporan-penjualan.png",
             caption="Bukti 8. Halaman laporan penjualan pada server demo setelah transaksi volume.",
             actors=["Supervisor Apotik", "Procurement", "Keuangan", "Auditor"],
             usecases=["Pilih periode dan filter", "Tampilkan penjualan/kulakan", "Telusuri rincian sumber", "Ekspor dan rekonsiliasi"],
             flow=["Pilih laporan", "Atur periode", "Terapkan filter", "Tampilkan agregat", "Drill-down dokumen", "Ekspor/rekonsiliasi"],
             gate="GATE: periode, toko, status dokumen, retur, diskon, pajak, dan pembatalan diperlakukan konsisten.",
             entities=["PENJUALAN\nheader/detail", "PENGADAAN\nPR-PO-BAST", "FILTER\nperiode/toko", "AGREGAT\nqty/nilai", "EKSPOR\nrekonsiliasi"],
             observation="endpoint laporan penjualan mengembalikan status success tetapi nol baris setelah 300 transaksi UAT; hal ini dicatat sebagai defect data/reporting, bukan PASS angka. Dashboard PR/PO/BAST/tagihan berhasil menunjukkan data pengadaan volume dan katalog laporan dapat dibuka",
             status="PASS akses layar / BLOCKED validasi isi laporan penjualan", controls="cut-off, status final, void/retur, diskon, pajak, COGS, vendor, toko, pagination, total, drill-down, dan konsistensi ekspor",
             acceptance="300 transaksi muncul pada periode yang tepat dan totalnya cocok dengan detail; 100 rantai pengadaan dapat ditelusuri; agregat layar, ekspor, dan dokumen sumber mempunyai nilai yang sama"),
        dict(title="Laporan Keuangan", screenshot=ACC / "24-laporan-laba-rugi.png",
             caption="Bukti 9. Laporan Laba Rugi berbasis jurnal pada server demo.",
             actors=["Akuntan", "Manajemen", "Auditor", "Owner COA"],
             usecases=["Tampilkan buku besar/trial balance", "Susun Laba Rugi dan Neraca", "Analisis Arus Kas", "Bandingkan periode dan sign-off"],
             flow=["Posting jurnal", "Buku besar", "Neraca saldo", "Laporan utama", "Bandingkan periode", "Review & sign-off"],
             gate="GATE: hanya jurnal posted, periode dan klasifikasi akun benar, saldo awal dan closing tervalidasi.",
             entities=["AKUN\nkode, grup", "DETAIL POSTED\nD/K", "BUKU BESAR\nsaldo", "TRIAL BALANCE\nD/K", "LAPORAN\nLR/neraca/arus kas"],
             observation="enam halaman inti—Laba Rugi, Neraca, Arus Kas, Keseluruhan Jurnal, Buku Besar, dan Neraca Saldo—berhasil dibuka dan dirender. Beberapa hasil belum mempunyai data yang dapat direkonsiliasi; Saldo Awal dan Jurnal Penyesuaian menampilkan masalah parsing respons, sedangkan katalog menyediakan laporan perbandingan namun belum diuji angka bulan/tahun",
             status="PASS render / BLOCKED rekonsiliasi angka dan beberapa backend", controls="klasifikasi akun, saldo awal, cut-off, posting, closing, retained earnings, arus kas, dimensi unit, perbandingan bulanan/tahunan, dan audit trail",
             acceptance="buku besar, trial balance, Laba Rugi, Neraca, dan Arus Kas seimbang serta cocok dengan 100 jurnal vendor dan transaksi sumber; perbandingan bulanan/tahunan dapat direproduksi dan disetujui owner Akuntansi"),
    ]


def build_document() -> Path:
    data = items()
    diagrams = make_diagrams(data)
    doc = Document()
    configure(doc)
    cover(doc)

    doc.add_heading("Ringkasan eksekutif", level=1)
    doc.add_paragraph(
        "UAT v1.34.23 dijalankan terhadap server demo.ecampus.id/ecampus menggunakan data bertanda rilis dan operasi idempoten. Run menghasilkan 100 transaksi obat jadi, 100 transaksi racikan, 100 transaksi gabungan, 100 rangkaian pengadaan pada setiap tahap PR–PO–BAST–tagihan–pembayaran, serta 100 jurnal pembayaran vendor yang terposting dan diverifikasi. Navigasi seluruh submenu Akuntansi dan enam laporan inti berhasil. Keputusan keseluruhan adalah LULUS BERSYARAT untuk UAT/pilot terkontrol. Produksi belum boleh dinyatakan siap karena katalog server hanya dua item, endpoint antrean farmasi belum terdeploy, laporan penjualan mengembalikan nol baris, beberapa jalur posting otomatis/saldo awal belum menghasilkan data yang dapat direkonsiliasi, dan artefak belum ditandatangani dengan keystore/sertifikat produksi."
    )
    note(doc, "Makna keputusan", "Keberhasilan transaksi dan jurnal tidak menghapus blocker deployment. Setiap blocker mempunyai kriteria retest di bagian temuan dan checklist rilis.", AMBER)

    doc.add_heading("Lingkungan, metode, dan volume", level=2)
    add_table(doc, ["Aspek", "Nilai / hasil"], [
        ("Server", "https://demo.ecampus.id/ecampus/ — tenant demo bersama"),
        ("Klien", "Flutter POS Apotik Windows; bukti 1600 × 900; API HTTPS"),
        ("Versi target", VERSION),
        ("Penjualan", "100 obat jadi + 100 racikan + 100 gabungan = 300 PASS"),
        ("Pengadaan", "100 PR + 100 PO + 100 BAST + 100 tagihan + 100 pembayaran disetujui"),
        ("Akuntansi", "100 jurnal vendor dibuat, posted, dan diverifikasi"),
        ("Master resep", "100 dibaca per page; summary provision server menyatakan 5.000 tersedia"),
        ("Master katalog", "2 item terbaca — blocker terhadap target data 100+"),
        ("Laporan", "6 laporan keuangan render; laporan penjualan success/0 rows — blocker rekonsiliasi"),
    ], [1.6, 4.8])

    doc.add_heading("Matriks keputusan", level=2)
    add_table(doc, ["No.", "Skenario", "Volume", "Status", "Temuan utama"], [
        (1, "Obat jadi", "100", "PASS", "Katalog volume belum terpenuhi"),
        (2, "Obat racik", "100", "PASS", "100 resep terbaca"),
        (3, "Gabungan", "100", "PASS transaksi", "Antrean live BLOCKED"),
        (4, "PR–pembayaran", "100/tahap", "PASS bersyarat", "Filter pembayaran marker 0"),
        (5, "Stok", "2 item", "BLOCKED volume", "Perlu seed/deploy katalog 100+"),
        (6, "Kedaluwarsa", "≥1 batch", "PASS monitor", "SOP disposisi perlu sign-off"),
        (7, "Jurnal vendor", "100", "PASS", "Posting otomatis lain perlu retest"),
        (8, "Laporan jual/kulak", "300/100", "BLOCKED angka", "Penjualan success tetapi 0 rows"),
        (9, "Laporan keuangan", "6 halaman", "PASS render", "Rekonsiliasi angka belum lulus"),
    ], [0.35, 1.6, 0.8, 1.05, 2.6])

    doc.add_heading("Pemetaan akun dari workbook pengguna", level=2)
    doc.add_paragraph("Workbook cetak_data_260904124814.xlsx berisi 317 akun data tanpa duplikasi kode pada pembacaan UAT. Mapping berikut dipakai sebagai rujukan, bukan perubahan diam-diam terhadap kebijakan akuntansi pemilik.")
    add_table(doc, ["Kode", "Nama akun", "Penggunaan"], [
        ("111.101", "KAS YAYASAN", "Kredit pembayaran vendor UAT"),
        ("151.200", "PERSEDIAAN BARANG LAINNYA", "Persediaan bila belum ada subakun obat khusus"),
        ("171.200", "UANG MUKA PEMBELIAN", "Termin/uang muka"),
        ("310.500", "HUTANG VENDOR", "Debit pelunasan vendor UAT"),
        ("310.600", "UTANG USAHA TOKO", "Alternatif setelah keputusan owner"),
        ("310.301", "HUTANG PPN", "Kewajiban PPN sesuai substansi"),
        ("410.900", "PENDAPATAN PENJUALAN TOKO", "Pendapatan Apotik"),
        ("510.900", "BEBAN POKOK PENJUALAN TOKO", "HPP penjualan"),
    ], [0.9, 2.7, 2.75])
    note(doc, "Kualitas data COA", "Kode 112,102 dan 121,109 memakai koma. Sistem tidak menormalisasikannya otomatis; owner COA harus menetapkan kode sah sebelum impor produksi.", AMBER)

    for idx, item in enumerate(data, 1):
        doc.add_page_break()
        doc.add_heading(f"{idx}. {item['title']}", level=1)
        add_picture(doc, item["screenshot"], item["caption"])
        doc.add_heading("Narasi UAT dan petunjuk penggunaan", level=2)
        paragraphs = narrative(item)
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
        note(doc, "Kelengkapan narasi", f"{word_count(paragraphs)} kata untuk screenshot ini; batas minimum yang diminta adalah 500 kata.", GREEN)

        doc.add_heading("Use case", level=2)
        add_picture(doc, diagrams[idx][0], f"Diagram {idx}.A — use case {item['title']}")
        doc.add_heading("Flowchart", level=2)
        add_picture(doc, diagrams[idx][1], f"Diagram {idx}.B — alur kerja {item['title']}")
        doc.add_heading("ERD dan aliran data", level=2)
        add_picture(doc, diagrams[idx][2], f"Diagram {idx}.C — model konseptual dan aliran data {item['title']}")

        doc.add_heading("Checklist operator dan acceptance", level=2)
        for step in item["flow"]:
            add_number(doc, step)
        add_table(doc, ["Kontrol", "Bukti yang harus ada", "Keputusan saat gagal"], [
            ("Otorisasi", "Role/menu dan approval", "Hentikan; perbaiki RBAC"),
            ("Integritas data", "Kode, referensi, detail, status", "Cari kode idempoten; jangan duplikasi"),
            ("Rekonsiliasi", item["acceptance"], "Tahan sign-off dan buat retest"),
        ], [1.15, 3.25, 2.0])

    doc.add_page_break()
    doc.add_heading("Temuan, risiko, dan retest wajib", level=1)
    add_table(doc, ["ID", "Temuan", "Dampak", "Retest / exit criterion", "Prioritas"], [
        ("APT-01", "Katalog hanya 2 item", "Cakupan kategori/stok belum representatif", "Deploy seed dan buktikan ≥100 item lintas kategori", "P0 sebelum produksi"),
        ("APT-02", "Endpoint antrean farmasi belum tersedia", "Layar publik belum live", "50 antrean live, status berubah, identitas tersamar", "P0 sebelum produksi"),
        ("APT-03", "Laporan penjualan success tetapi 0 rows", "Rekonsiliasi pendapatan gagal", "300 transaksi muncul dan total cocok", "P0 sebelum produksi"),
        ("APT-04", "Filter daftar pembayaran marker 0", "Pencarian/audit dokumen sulit", "100 pembayaran ditemukan melalui filter yang sama", "P1"),
        ("APT-05", "Saldo Awal/Jurnal Penyesuaian parse error", "Opening/adjustment belum dapat divalidasi", "Respons valid dan jurnal seimbang", "P0 sebelum closing"),
        ("APT-06", "Posting otomatis tertentu kosong/belum siap", "Jurnal memerlukan prosedur manual", "Preview, mapping, posting, rollback terverifikasi", "P0 sebelum otomatisasi"),
        ("APT-07", "RenderFlex overflow 34 px pada katalog/anggaran", "Tampilan terpotong pada resolusi tertentu", "Tidak ada overflow pada 1366×768 dan 1600×900", "P2"),
        ("REL-01", "APK debug-signed; EXE tanpa Authenticode", "Tidak layak distribusi produksi", "Tanda tangan resmi dan verifikasi publisher/signature", "P0 produksi"),
    ], [0.6, 1.65, 1.55, 2.25, 0.8])

    doc.add_heading("Checklist pra-rilis dan deployment", level=1)
    for text in [
        "Semua tes fokus, build Android, build Windows, dan kompilasi installer selesai tanpa error.",
        "Nama varian, application ID, versionName/versionCode, file version, dan tag GitHub sama dengan v1.34.23/build 185.",
        "APK dan installer EXE diverifikasi checksum SHA-256; aset Windows tidak dipublikasikan sebagai ZIP.",
        "Tidak ada kredensial, token, data pasien nyata, atau file key signing di commit maupun release asset.",
        "Backup, rollback, owner, dan jendela deployment ditetapkan sebelum deploy backend atau migrasi data.",
        "Retest endpoint antrean, katalog, laporan, posting otomatis, saldo awal, dan perbandingan periode dijadwalkan setelah backend sesuai terdeploy.",
        "Stakeholder memahami bahwa release saat ini ditandai UAT/pilot karena signing produksi belum tersedia.",
    ]:
        add_bullet(doc, f"☐ {text}")

    doc.add_heading("Rollback dan penghentian pilot", level=2)
    doc.add_paragraph("Hentikan pilot bila transaksi gagal tersimpan, stok menjadi negatif/tidak dapat ditelusuri, batch kedaluwarsa dapat dijual, identitas pasien terbuka pada layar publik, dokumen pengadaan terduplikasi, pembayaran tidak memiliki sumber sah, jurnal tidak seimbang, atau laporan material berbeda dari dokumen sumber. Rollback klien dilakukan dengan memasang kembali artefak rilis sebelumnya yang telah diverifikasi checksum. Perubahan backend/data harus memakai backup dan skrip reversal yang disetujui pemilik sistem; jangan menghapus transaksi audit secara manual.")

    doc.add_heading("Sign-off", level=1)
    add_table(doc, ["Peran", "Nama", "Keputusan", "Tanggal / tanda tangan"], [
        ("Owner Apotik", "", "", ""),
        ("Apoteker Penanggung Jawab", "", "", ""),
        ("Owner Procurement", "", "", ""),
        ("Owner Keuangan/Akuntansi", "", "", ""),
        ("IT/DevOps", "", "", ""),
        ("Keamanan/Privasi", "", "", ""),
    ], [1.8, 1.6, 1.5, 1.5])

    doc.add_heading("Lampiran: artefak mesin", level=1)
    for path in [SHOT / "uat-summary.json", SHOT / "procurement-summary.json", SHOT / "vendor-journal-summary.json"]:
        add_bullet(doc, f"{path.relative_to(ROOT)}")
    doc.add_paragraph("Dokumen ini hanya menampilkan sembilan screenshot representatif agar setiap gambar mempunyai narasi minimal 500 kata serta diagram use case, flowchart, dan ERD/aliran data. Seluruh screenshot mentah tetap berada di folder UAT v1.34.23 pada repository.")

    output = OUT / "Laporan-UAT-dan-Manual-POS-Apotik-v1.34.23.docx"
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    manifest = {
        "version": VERSION,
        "document": output.name,
        "screenshotSections": len(data),
        "narrativeWordCounts": {item["title"]: word_count(narrative(item)) for item in data},
        "diagramCount": len(diagrams) * 3,
    }
    (OUT / "document-manifest.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")
    presentation_data = []
    for idx, item in enumerate(data, 1):
        presentation_data.append({
            "index": idx,
            "title": item["title"],
            "screenshot": item["screenshot"].relative_to(ROOT).as_posix(),
            "caption": item["caption"],
            "actors": item["actors"],
            "usecases": item["usecases"],
            "flow": item["flow"],
            "gate": item["gate"],
            "entities": item["entities"],
            "observation": item["observation"],
            "status": item["status"],
            "controls": item["controls"],
            "acceptance": item["acceptance"],
            "narrative": narrative(item),
            "narrativeWordCount": word_count(narrative(item)),
        })
    (OUT / "presentation-data.json").write_text(
        json.dumps(presentation_data, indent=2, ensure_ascii=False), encoding="utf-8")
    return output


if __name__ == "__main__":
    print(build_document())
