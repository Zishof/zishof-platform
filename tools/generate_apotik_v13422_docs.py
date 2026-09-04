from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "pos-apotik-emedik" / "uat-v1.34.22"
SHOT = OUT / "screenshots"
PROC = OUT / "screenshots-pengadaan"
ACC = OUT / "screenshots-akuntansi"
DIAGRAM = OUT / "diagrams"
VERSION = "1.34.22 (build 184)"
DATE = "4 September 2026"

NAVY = "102238"
BLUE = "075985"
TEAL = "0F766E"
GREEN = "15803D"
AMBER = "B45309"
RED = "B91C1C"
LIGHT = "F3F7FA"
MID = "DCE8F0"
INK = "172033"


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def rounded(draw, xy, fill, outline="#D4E2EA", radius=22, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw, xy, text, fnt, fill="#172033", max_chars=24, spacing=8):
    lines = []
    for paragraph in text.split("\n"):
        lines.extend(wrap(paragraph, max_chars) or [""])
    y = xy[1] + 14
    for line in lines:
        box = draw.textbbox((0, 0), line, font=fnt)
        x = (xy[0] + xy[2] - (box[2] - box[0])) / 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += box[3] - box[1] + spacing


def arrow(draw, start, end, color="#075985", width=5):
    draw.line([start, end], fill=color, width=width)
    x, y = end
    if abs(end[0] - start[0]) >= abs(end[1] - start[1]):
        d = 1 if end[0] > start[0] else -1
        draw.polygon([(x, y), (x - d * 18, y - 10), (x - d * 18, y + 10)], fill=color)
    else:
        d = 1 if end[1] > start[1] else -1
        draw.polygon([(x, y), (x - 10, y - d * 18), (x + 10, y - d * 18)], fill=color)


def save_diagram(name, title, boxes, arrows, notes=None):
    DIAGRAM.mkdir(parents=True, exist_ok=True)
    im = Image.new("RGB", (1600, 900), "#F7FAFC")
    draw = ImageDraw.Draw(im)
    draw.rectangle((0, 0, 1600, 92), fill="#102238")
    draw.text((48, 25), title, font=font(34, True), fill="white")
    for box in boxes:
        xy, label, fill, outline = box
        rounded(draw, xy, fill, outline, radius=20, width=3)
        center_text(draw, xy, label, font(22, True), max_chars=24)
    for start, end, color in arrows:
        arrow(draw, start, end, color=color)
    if notes:
        y = 800
        for note in notes:
            draw.text((55, y), f"• {note}", font=font(18), fill="#334155")
            y += 28
    path = DIAGRAM / name
    im.save(path, quality=96)
    return path


def diagrams():
    files = {}
    files["usecase"] = save_diagram(
        "01-use-case-pos-apotik.png",
        "Use Case POS Apotik — Peran, Kontrol, dan Penerima Informasi",
        [
            ((40, 150, 260, 245), "Kasir", "#DBEAFE", "#2563EB"),
            ((40, 330, 260, 425), "Apoteker", "#EDE9FE", "#7C3AED"),
            ((40, 510, 260, 605), "Gudang / Penerimaan", "#FEF3C7", "#D97706"),
            ((1340, 150, 1560, 245), "Procurement", "#E0F2FE", "#0284C7"),
            ((1340, 330, 1560, 425), "Keuangan / Akuntansi", "#DCFCE7", "#16A34A"),
            ((1340, 510, 1560, 605), "Pasien / Keluarga", "#FCE7F3", "#DB2777"),
            ((460, 135, 1140, 235), "Kelola formularium, batch, FEFO, dan kedaluwarsa", "white", "#0F766E"),
            ((460, 270, 1140, 370), "Jual obat jadi • telaah & siapkan obat racikan", "white", "#0F766E"),
            ((460, 405, 1140, 505), "PR → PO → BAST → tagihan → pembayaran vendor", "white", "#0F766E"),
            ((460, 540, 1140, 640), "Jurnal, posting, laporan keuangan, dan audit trail", "white", "#0F766E"),
            ((460, 675, 1140, 765), "Layar antrean publik dengan identitas tersamar", "white", "#0F766E"),
        ],
        [
            ((260, 198), (460, 185), "#2563EB"), ((260, 378), (460, 320), "#7C3AED"),
            ((260, 558), (460, 455), "#D97706"), ((1340, 198), (1140, 455), "#0284C7"),
            ((1340, 378), (1140, 590), "#16A34A"), ((1340, 558), (1140, 720), "#DB2777"),
        ],
        ["Hak menu mengikuti RBAC; pengguna hanya melihat dan menjalankan kewenangannya.",
         "Identitas lengkap pasien tetap berada di layar petugas, bukan layar publik."],
    )
    files["finished"] = save_diagram(
        "02-flow-obat-jadi.png", "Flow Penjualan Obat Jadi — Dari Permintaan sampai Serah Terima",
        [
            ((45, 290, 245, 410), "1. Identifikasi pasien / pembeli", "#DBEAFE", "#2563EB"),
            ((300, 290, 500, 410), "2. Cari / pindai obat", "white", "#0284C7"),
            ((555, 290, 755, 410), "3. Validasi stok, batch, ED, golongan", "white", "#0F766E"),
            ((810, 290, 1010, 410), "4. FEFO & verifikasi jumlah", "white", "#0F766E"),
            ((1065, 290, 1265, 410), "5. Bayar & catat transaksi", "#DCFCE7", "#16A34A"),
            ((1320, 290, 1555, 410), "6. Etiket, edukasi, serah terima", "#EDE9FE", "#7C3AED"),
            ((555, 550, 1010, 670), "Jika obat terkendali/LASA: konfirmasi ganda dan dokumentasikan alasan", "#FEF3C7", "#D97706"),
        ],
        [((245, 350), (300, 350), "#075985"), ((500, 350), (555, 350), "#075985"),
         ((755, 350), (810, 350), "#075985"), ((1010, 350), (1065, 350), "#075985"),
         ((1265, 350), (1320, 350), "#075985"), ((655, 410), (655, 550), "#D97706"),
         ((910, 550), (910, 410), "#D97706")],
        ["Jangan menjual batch kedaluwarsa; pilih batch dengan tanggal kedaluwarsa terdekat yang masih layak.",
         "Serah terima wajib mencocokkan pasien, nama obat, jumlah, etiket, dan instruksi."],
    )
    files["compound"] = save_diagram(
        "03-flow-obat-racikan.png", "Flow Resep dan Obat Racikan — Telaah Klinis, Produksi, dan Penyerahan",
        [
            ((55, 180, 305, 290), "Resep masuk / dipilih", "#DBEAFE", "#2563EB"),
            ((375, 180, 625, 290), "Telaah administratif & farmasetik", "white", "#7C3AED"),
            ((695, 180, 945, 290), "Klarifikasi dokter bila perlu", "#FEF3C7", "#D97706"),
            ((1015, 180, 1265, 290), "Hitung bahan & pilih batch FEFO", "white", "#0F766E"),
            ((1335, 180, 1545, 290), "Racik, kemas, etiket", "#EDE9FE", "#7C3AED"),
            ((1335, 480, 1545, 590), "Serahkan & edukasi", "#DCFCE7", "#16A34A"),
            ((1015, 480, 1265, 590), "Pemeriksaan akhir oleh petugas kedua", "white", "#0F766E"),
            ((695, 480, 945, 590), "Status SIAP + panggil antrean", "#DBEAFE", "#2563EB"),
            ((375, 480, 625, 590), "Rekonsiliasi stok & audit trail", "white", "#334155"),
        ],
        [((305, 235), (375, 235), "#075985"), ((625, 235), (695, 235), "#075985"),
         ((945, 235), (1015, 235), "#075985"), ((1265, 235), (1335, 235), "#075985"),
         ((1440, 290), (1440, 480), "#075985"), ((1335, 535), (1265, 535), "#075985"),
         ((1015, 535), (945, 535), "#075985"), ((695, 535), (625, 535), "#075985")],
        ["Pisahkan area racik, alat, dan bahan; catat penyimpangan serta tindakan koreksi.",
         "Layar publik hanya menampilkan nomor antrean dan identitas tersamar."],
    )
    files["p2p"] = save_diagram(
        "04-flow-procure-to-pay.png", "Procure-to-Pay Obat — PR sampai Pembayaran Vendor",
        [
            ((45, 260, 245, 380), "PR\nKebutuhan & justifikasi", "#DBEAFE", "#2563EB"),
            ((305, 260, 505, 380), "Persetujuan\nAnggaran & otorisasi", "white", "#0284C7"),
            ((565, 260, 765, 380), "PO\nTermin / non-termin", "white", "#0F766E"),
            ((825, 260, 1025, 380), "BAST\nQty, mutu, batch, ED", "#FEF3C7", "#D97706"),
            ((1085, 260, 1285, 380), "Terima Tagihan\n3-way match", "white", "#7C3AED"),
            ((1345, 260, 1555, 380), "Bayar Vendor\nTransfer & bukti", "#DCFCE7", "#16A34A"),
            ((825, 520, 1285, 650), "Kontrol: PR–PO–BAST–faktur harus konsisten; selisih ditahan dan ditindaklanjuti", "#FEE2E2", "#B91C1C"),
        ],
        [((245, 320), (305, 320), "#075985"), ((505, 320), (565, 320), "#075985"),
         ((765, 320), (825, 320), "#075985"), ((1025, 320), (1085, 320), "#075985"),
         ((1285, 320), (1345, 320), "#075985"), ((925, 380), (925, 520), "#B91C1C"),
         ((1185, 520), (1185, 380), "#B91C1C")],
        ["UAT volume: 50 PR, 50 PO, 50 BAST, 50 tagihan, dan 50 pembayaran disetujui.",
         "Satu proses transfer berisi 50 detail pembayaran dengan total Rp45.730.000."],
    )
    files["accounting"] = save_diagram(
        "05-flow-akuntansi.png", "Flow Akuntansi Pembayaran Vendor — Dokumen sampai Laporan",
        [
            ((65, 180, 335, 300), "Pembayaran vendor disetujui", "#DBEAFE", "#2563EB"),
            ((405, 180, 675, 300), "Preview jurnal otomatis", "#FEF3C7", "#D97706"),
            ((745, 180, 1015, 300), "Jika belum tersedia: Jurnal Umum manual", "#FEE2E2", "#B91C1C"),
            ((1085, 180, 1355, 300), "Validasi debit = kredit", "white", "#0F766E"),
            ((1085, 480, 1355, 600), "Posting & kunci bukti", "#DCFCE7", "#16A34A"),
            ((745, 480, 1015, 600), "Buku Besar / Neraca Saldo", "white", "#0F766E"),
            ((405, 480, 675, 600), "Laba Rugi / Neraca / Arus Kas", "#EDE9FE", "#7C3AED"),
            ((65, 480, 335, 600), "Rekonsiliasi dan sign-off", "white", "#334155"),
        ],
        [((335, 240), (405, 240), "#075985"), ((675, 240), (745, 240), "#B45309"),
         ((1015, 240), (1085, 240), "#075985"), ((1220, 300), (1220, 480), "#075985"),
         ((1085, 540), (1015, 540), "#075985"), ((745, 540), (675, 540), "#075985"),
         ((405, 540), (335, 540), "#075985")],
        ["Jurnal UAT: Dr 310.500 HUTANG VENDOR / Cr 111.101 KAS YAYASAN.",
         "50 jurnal manual terposting; preview otomatis pembayaran vendor belum tersedia pada server demo."],
    )
    return files


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=INK, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, True, "FFFFFF", 8.5)
        shade(t.rows[0].cells[i], BLUE)
        if widths:
            t.rows[0].cells[i].width = Cm(widths[i])
    set_repeat_table_header(t.rows[0])
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, INK, 8)
            if ridx % 2:
                shade(cells[i], "F8FAFC")
            if widths:
                cells[i].width = Cm(widths[i])
    return t


def configure(doc, title, kind):
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.6)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in [("Title", 28, NAVY), ("Heading 1", 19, NAVY),
                              ("Heading 2", 14, BLUE), ("Heading 3", 11, TEAL)]:
        s = styles[name]
        s.font.name = "Aptos Display"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(5)
    header = sec.header.paragraphs[0]
    header.text = f"POS APOTIK • {kind.upper()} • v{VERSION}"
    header.style = styles["Caption"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Dokumen terkendali • ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)
    props = doc.core_properties
    props.title = title
    props.subject = f"POS Apotik {VERSION}"
    props.author = "Tim eBisnis / Zishof"
    props.keywords = "POS Apotik, UAT, user manual, farmasi, pengadaan, akuntansi"


def cover(doc, eyebrow, title, subtitle, status, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(95)
    r = p.add_run(eyebrow.upper())
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(color)
    p = doc.add_paragraph(style="Title")
    p.add_run(title)
    p = doc.add_paragraph()
    r = p.add_run(subtitle)
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string("475569")
    p.paragraph_format.space_after = Pt(24)
    t = doc.add_table(rows=4, cols=2)
    t.style = "Table Grid"
    data = [("Versi", VERSION), ("Tanggal", DATE), ("Status", status),
            ("Cakupan", "Desktop Windows, Android, API AIS, pengadaan, akuntansi, layar publik")]
    for i, (k, v) in enumerate(data):
        set_cell_text(t.cell(i, 0), k, True, "FFFFFF")
        shade(t.cell(i, 0), color)
        set_cell_text(t.cell(i, 1), v)
    doc.add_paragraph()
    p = doc.add_paragraph("Kerahasiaan: internal implementasi dan penerimaan pengguna. Bukti layar pasien memakai identitas tersamar.")
    p.style = doc.styles["Caption"]
    doc.add_page_break()


def note(doc, label, text, color=BLUE):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.cell(0, 0)
    shade(c, "EAF3F8")
    p = c.paragraphs[0]
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def screenshot(doc, path, caption, width=6.85):
    if not path.exists():
        note(doc, "Bukti tidak tersedia", str(path), RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    p = doc.add_paragraph(caption)
    p.style = doc.styles["Caption"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def section_text(doc, title, intro, steps, controls, image=None, caption=None):
    doc.add_heading(title, level=1)
    doc.add_paragraph(intro)
    doc.add_heading("Prosedur kerja", level=2)
    for step in steps:
        numbered(doc, step)
    doc.add_heading("Kontrol dan hasil yang diharapkan", level=2)
    for item in controls:
        bullet(doc, item)
    if image:
        screenshot(doc, image, caption or title)


def manual(diag):
    doc = Document()
    configure(doc, "Manual Pengguna POS Apotik", "Manual Pengguna")
    cover(doc, "Panduan operasional terkendali", "Manual Pengguna POS Apotik",
          "Dari pelayanan obat jadi dan racikan sampai pengadaan, pembayaran vendor, jurnal, dan laporan",
          "Siap digunakan dengan catatan prasyarat server pada Bab 20")

    doc.add_heading("Cara menggunakan manual ini", level=1)
    doc.add_paragraph(
        "Manual ini disusun untuk kasir, tenaga teknis kefarmasian, apoteker, petugas gudang, procurement, "
        "keuangan, akuntansi, supervisor, dan administrator aplikasi. Setiap bab menjelaskan tujuan menu, "
        "prasyarat, urutan tindakan, titik kontrol, hasil yang benar, serta cara menangani penyimpangan. "
        "Istilah tombol mengikuti tampilan versi 1.34.22. Bila nama kebijakan lokal berbeda, institusi wajib "
        "menetapkan padanannya dalam SOP tanpa menghilangkan kontrol keselamatan pasien, segregasi tugas, "
        "jejak audit, dan rekonsiliasi keuangan."
    )
    note(doc, "Prinsip keselamatan", "Aplikasi mendukung proses, tetapi keputusan klinis tetap dilakukan apoteker/tenaga kesehatan yang berwenang. Jangan menjadikan data demo sebagai dasar terapi pasien nyata.")
    doc.add_heading("Daftar bab", level=2)
    chapters = [
        "1. Ruang lingkup dan pembagian peran", "2. Instalasi Windows dan Android",
        "3. Alamat server, login, dan RBAC", "4. Dashboard dan pembukaan operasional",
        "5. Formularium dan profil obat", "6. Batch, FEFO, kedaluwarsa, dan stok",
        "7. Penerimaan PBF", "8. Penjualan obat jadi", "9. Resep dan obat racikan",
        "10. Layar pasien/keluarga dan multi-monitor", "11. Permintaan Pembelian (PR)",
        "12. Pemesanan Pembelian (PO)", "13. BAST penerimaan barang/jasa",
        "14. Terima Tagihan Vendor", "15. Pembayaran vendor dan Proses Transfer",
        "16. Pemetaan akun dan Jurnal Umum", "17. Posting dan laporan keuangan",
        "18. Rekonsiliasi, tutup hari, dan audit", "19. Bantuan dan penanganan masalah",
        "20. Batasan versi, prasyarat deploy, dan checklist go-live", "Lampiran A. Akun rujukan",
        "Lampiran B. Kamus status dan checklist cepat"
    ]
    for c in chapters:
        bullet(doc, c)
    doc.add_page_break()

    doc.add_heading("1. Ruang lingkup dan pembagian peran", level=1)
    doc.add_paragraph(
        "POS Apotik menyatukan pelayanan farmasi, pengelolaan persediaan berbasis batch, pengadaan, utang "
        "vendor, jurnal, dan laporan. Pemisahan peran penting karena orang yang meminta barang sebaiknya tidak "
        "menyetujui sendiri, penerima barang harus membandingkan dengan PO, petugas pembayaran harus bekerja "
        "dari tagihan yang telah disetujui, dan akuntansi harus merekonsiliasi dokumen sumber dengan jurnal."
    )
    table(doc, ["Peran", "Tanggung jawab utama", "Tidak boleh diabaikan"], [
        ("Kasir", "Identifikasi pembeli, pilih obat/batch, pembayaran, bukti transaksi", "Tidak menjual obat kedaluwarsa atau melompati konfirmasi obat terkendali"),
        ("TTK/Apoteker", "Telaah resep, racik, periksa akhir, edukasi, ubah status antrean", "Klarifikasi resep tidak jelas dan dokumentasikan intervensi"),
        ("Gudang", "Terima PBF, input batch/ED, inspeksi qty dan mutu", "Karantina selisih, rusak, atau suhu tidak sesuai"),
        ("Procurement", "PR, PO, vendor, termin, pemantauan pemenuhan", "Pastikan otorisasi dan tiga-way match"),
        ("Keuangan", "Terima tagihan, jadwalkan/bayar, bukti transfer", "Bayar hanya tagihan sah dan terotorisasi"),
        ("Akuntansi", "Pemetaan akun, jurnal, posting, rekonsiliasi, laporan", "Debit harus sama dengan kredit; bukti wajib tertaut"),
        ("Admin", "Server, role/menu, konfigurasi, deployment, backup", "Hak minimum, audit perubahan, dan data-sample nonaktif di produksi"),
    ], [3.0, 7.0, 7.0])
    screenshot(doc, diag["usecase"], "Gambar 1. Use case lintas peran POS Apotik")

    section_text(doc, "2. Instalasi Windows dan Android",
        "Gunakan hanya paket dari GitHub Release resmi. Cocokkan versi, nama varian Apotik, ukuran berkas, dan SHA-256 pada catatan rilis sebelum instalasi. APK versi ini dapat dibangun dengan debug key bila keystore produksi belum disediakan; kondisi penandatanganan harus diperiksa sebelum distribusi produksi.",
        [
            "Windows: unduh installer/arsip desktop Apotik, verifikasi SHA-256, tutup aplikasi lama, lalu jalankan installer sebagai pengguna yang berhak memasang aplikasi.",
            "Android: aktifkan pemasangan dari sumber yang dipercaya hanya selama instalasi, pilih APK varian Apotik, periksa nama paket dan versi, kemudian nonaktifkan kembali izin sumber tidak dikenal.",
            "Buka aplikasi, pastikan branding menunjukkan POS Apotik, bukan eMedik atau varian generik.",
            "Masukkan host dan context path yang diberikan administrator. Gunakan HTTPS; jangan menonaktifkan validasi sertifikat.",
            "Login dengan akun uji lebih dahulu, lakukan sinkronisasi, lalu uji satu transaksi nonproduksi sebelum go-live.",
        ],
        ["Versi tampilan harus 1.34.22 build 184.", "Backup basis data lokal sebelum upgrade mayor.",
         "Simpan installer lama sampai smoke test versi baru selesai agar rollback tersedia."],
    )

    section_text(doc, "3. Alamat server, login, dan RBAC",
        "Akun menentukan menu dan aksi yang dapat dilakukan. Pengalihan setelah login dapat diarahkan langsung ke halaman Apotik/JSP oleh konfigurasi role. Bila menu hilang, jangan meminjam akun lain; minta admin memeriksa role, izin create/read/update/approve/post, toko aktif, dan tenant.",
        [
            "Di Pengaturan Server masukkan host tanpa protokol, context path aplikasi, dan aktifkan HTTPS.",
            "Uji koneksi. Bila gagal, catat waktu, host, jaringan, dan kode referensi yang tampil.",
            "Masukkan nama pengguna dan kata sandi sendiri. Jangan menyimpan kata sandi di dokumen atau mengirimkannya melalui pesan umum.",
            "Setelah login, periksa nama pengguna, peran, toko/instalasi farmasi, tenant, dan versi di sidebar.",
            "Tekan Sinkronkan bila aplikasi menawarkan penyiapan data lokal. Tunggu selesai sebelum transaksi pertama.",
        ],
        ["RBAC memakai prinsip least privilege.", "Admin global tanpa toko tidak boleh dipakai sebagai kasir produksi.",
         "Logout setelah selesai; kunci perangkat ketika meninggalkan meja."],
        SHOT / "09-laporan-penjualan.png", "Gambar 2. Shell aplikasi menampilkan identitas pengguna, menu berbasis role, sinkronisasi, Bantuan, dan versi")

    section_text(doc, "4. Dashboard dan pembukaan operasional",
        "Dashboard menempatkan pekerjaan berisiko di bagian atas: resep menunggu, batch mendekati kedaluwarsa, stok habis, dan jumlah obat terbaca. Gunakan dashboard sebagai briefing awal shift dan daftar tindakan, bukan hanya statistik.",
        [
            "Login dan pastikan tanggal/jam perangkat benar.", "Periksa jumlah resep menunggu dan tentukan prioritas klinis.",
            "Buka batch mendekati kedaluwarsa, pisahkan yang sudah kedaluwarsa, dan rencanakan FEFO untuk yang masih layak.",
            "Periksa stok habis/minimum lalu buat kebutuhan PR bila pengadaan diperlukan.",
            "Konfirmasi printer, pemindai barcode, laci kas, jaringan, dan layar pasien berfungsi.",
            "Catat serah terima shift: antrean belum selesai, obat karantina, transaksi tertunda, dan selisih yang belum ditutup.",
        ],
        ["Resep tidak boleh dibiarkan tanpa pemilik proses.", "Batch kedaluwarsa tidak boleh berada di rak jual.",
         "Gangguan sinkronisasi harus diketahui supervisor sebelum operasi dilanjutkan."],
        SHOT / "00-dashboard-operasional.png", "Gambar 3. Dashboard operasional dengan 50 resep menunggu dan indikator batch")

    section_text(doc, "5. Formularium dan profil obat",
        "Formularium adalah sumber data nama, kode, golongan, harga, satuan, tanda LASA/high-alert, dan sifat terkendali. Kesalahan master menyebar ke pencarian kasir, stok, etiket, laporan, serta kontrol klinis; perubahan harus ditinjau dan dapat diaudit.",
        [
            "Cari obat berdasarkan nama generik, merek, kode, atau barcode.",
            "Periksa keunikan kode dan konsistensi satuan beli, satuan stok, serta satuan jual.",
            "Isi golongan obat dan tandai LASA/high-alert/terkendali sesuai kebijakan farmasi.",
            "Tetapkan harga beli/jual sesuai kewenangan; jangan menimpa histori transaksi lama.",
            "Simpan, cari ulang, lalu uji kemunculannya pada kasir dan laporan.",
        ],
        ["Gunakan penamaan yang membedakan kekuatan dan bentuk sediaan.", "Barcode tidak boleh dipakai oleh dua item aktif.",
         "Perubahan kritis memerlukan maker-checker."],
        SHOT / "03-formularium-obat.png", "Gambar 4. Formularium obat pada komponen produksi")

    section_text(doc, "6. Batch, FEFO, kedaluwarsa, dan stok",
        "Persediaan farmasi harus ditelusuri sampai batch dan tanggal kedaluwarsa. FEFO memilih tanggal kedaluwarsa paling dekat yang masih aman, bukan sekadar stok tertua. Obat kedaluwarsa, rusak, recall, atau mengalami deviasi suhu harus dikarantina dan tidak tersedia untuk penjualan.",
        [
            "Buka Obat & Persediaan lalu pilih Batch & Expiry.", "Filter mendekati kedaluwarsa sesuai horizon kebijakan.",
            "Bandingkan stok sistem dengan fisik dan lokasi.", "Untuk selisih, lakukan investigasi sebelum opname; simpan alasan dan bukti.",
            "Saat memilih batch di kasir/racikan, gunakan batch layak dengan ED terdekat.",
            "Review laporan kedaluwarsa dan register obat terkendali secara periodik.",
        ],
        ["Tidak ada stok negatif tanpa insiden yang terdokumentasi.", "Batch expired tidak dapat dipilih.",
         "Recall dapat ditelusuri ke penerimaan, transaksi, dan pasien sesuai kewenangan privasi."],
        SHOT / "04-batch-kedaluwarsa.png", "Gambar 5. Monitor batch dan kedaluwarsa")

    section_text(doc, "7. Penerimaan PBF",
        "Penerimaan PBF menambah persediaan dan, bila tanggal kedaluwarsa diisi, membentuk batch. Petugas harus memeriksa dokumen dan kondisi fisik sebelum menyimpan karena kesalahan penerimaan memengaruhi stok, nilai persediaan, BAST, utang, dan kesiapan pelayanan.",
        [
            "Siapkan PO/nota pengiriman, faktur bila sudah ada, dan bukti rantai dingin untuk item terkait.",
            "Pilih Penerimaan PBF; isi vendor, nomor faktur, lokasi, dan catatan.",
            "Tambahkan obat; masukkan kuantitas diterima, harga beli, batch, dan tanggal kedaluwarsa.",
            "Bandingkan jumlah, mutu kemasan, kekuatan, suhu, dan ED dengan pesanan.",
            "Karantina selisih/rusak; jangan mencampur dengan stok layak.",
            "Simpan lalu verifikasi batch muncul dan stok bertambah. Pada UAT dicatat penerimaan 150 unit UJI-PCT untuk mendukung 100 transaksi.",
        ],
        ["Nomor dokumen dan vendor wajib dapat ditelusuri.", "Tanggal kedaluwarsa memakai format ISO pada API dan format lokal di UI.",
         "Penerimaan langsung dan BAST pengadaan harus direkonsiliasi agar tidak menggandakan stok."],
        SHOT / "05-penerimaan-pbf.png", "Gambar 6. Form penerimaan PBF")

    section_text(doc, "8. Penjualan obat jadi",
        "Mode OTC/Obat Bebas digunakan untuk obat yang boleh dilayani tanpa resep sesuai peraturan dan kebijakan. Mode Resep Dokter dipakai bila transaksi terkait resep. Sistem menunjukkan label terkendali dan LASA; petugas wajib menindaklanjuti peringatan dan tidak sekadar mengeklik konfirmasi.",
        [
            "Pilih mode transaksi yang benar dan identifikasi pembeli/pasien.", "Cari nama/kode atau pindai barcode; pastikan kekuatan, bentuk, dan satuan.",
            "Tambahkan obat, periksa stok serta batch. Sistem harus menolak batch kedaluwarsa.",
            "Untuk LASA atau obat terkendali, lakukan pembacaan ulang dan verifikasi orang kedua sesuai SOP.",
            "Pilih metode pembayaran; cek total, uang diterima, kembalian, dan referensi nontunai.",
            "Selesaikan pembayaran, cetak/kirim bukti, pasang etiket, edukasi pasien, dan catat serah terima.",
        ],
        ["Nama pasien/pembeli, obat, jumlah, harga, batch, dan waktu transaksi harus terekam.",
         "Jangan menggabungkan transaksi pasien berbeda.", "UAT v1.34.22 berhasil memproses 50 transaksi obat jadi."],
        SHOT / "01-kasir-obat-jadi.png", "Gambar 7. Kasir obat jadi menampilkan stok, label LASA, dan obat terkendali")
    screenshot(doc, diag["finished"], "Gambar 8. Flow terkontrol penjualan obat jadi")

    section_text(doc, "9. Resep dan obat racikan",
        "Alur resep memerlukan telaah administratif, farmasetik, dan klinis sesuai kewenangan. Racikan menambah risiko perhitungan dosis, kompatibilitas, kontaminasi, ketepatan pembagian, pelabelan, dan stabilitas. Status antrean harus mencerminkan keadaan nyata agar keluarga pasien tidak dipanggil sebelum pemeriksaan akhir.",
        [
            "Buka daftar Resep; pilih resep menunggu dan cocokkan pasien, dokter, tanggal, serta item.",
            "Telaah dosis, frekuensi, duplikasi, alergi/interaksi yang tersedia, bentuk sediaan, dan keterbacaan instruksi.",
            "Jika meragukan, tahan proses dan dokumentasikan klarifikasi dokter.",
            "Hitung kebutuhan bahan serta jumlah bungkus/kapsul; minta verifikasi perhitungan.",
            "Ambil bahan dengan FEFO, catat batch, timbang/ukur, racik di area sesuai, lalu kemas.",
            "Buat etiket yang jelas: identitas, aturan pakai, penyimpanan, tanggal, dan beyond-use date sesuai kebijakan.",
            "Lakukan pemeriksaan akhir, ubah status SIAP, panggil antrean, cocokkan identitas lengkap secara privat, lalu edukasi.",
        ],
        ["UAT membaca 100 resep/racikan dan berhasil memproses 50 transaksi penebusan racikan.",
         "Identitas lengkap dan informasi klinis tidak ditampilkan di layar publik.",
         "Setiap penyimpangan harus memiliki alasan, penanggung jawab, dan tindak lanjut."],
        SHOT / "02-resep-racikan.png", "Gambar 9. Antrean resep menampilkan 50 resep menunggu pada halaman")
    screenshot(doc, diag["compound"], "Gambar 10. Flow resep dan racikan")

    section_text(doc, "10. Layar pasien/keluarga dan multi-monitor",
        "Layar farmasi membantu pasien/keluarga mengetahui obat jadi atau racikan yang menunggu, disiapkan, atau siap diambil. Desain memisahkan kolom obat jadi dan racikan, menyediakan panel edukasi, jam, loket, serta pita panggilan. Nama dan nomor rekam medis disamarkan; diagnosis, telepon, alamat, dan aturan pakai tidak diminta oleh layar publik.",
        [
            "Dari menu Layar Pelanggan/Farmasi pilih toko/instalasi dan mode Semua, Obat Jadi, atau Racikan.",
            "Pilih Buka di layar kedua. Pindahkan jendela ke monitor pasien dan aktifkan layar penuh.",
            "Untuk monitor tambahan, buka kembali mode khusus Obat Jadi atau Racikan; jumlah monitor dapat ditambah sesuai kemampuan perangkat.",
            "Pastikan kartu menampilkan nomor antrean, nama tersamar, RM tersamar, status, dan loket—tanpa data klinis sensitif.",
            "Ubah status pada konsol petugas sesuai progres nyata. Ketika SIAP, verifikasi panggilan muncul pada pita hijau.",
            "Jika koneksi terputus, layar mempertahankan data terakhir dan memberi indikator gangguan; petugas beralih ke panggilan manual.",
        ],
        ["Monitor harus menghadap area tunggu, bukan area publik yang tidak relevan.",
         "Jangan menyebut diagnosis atau detail obat sensitif melalui pengeras suara.",
         "Pada UAT ini UI layar lulus dengan 50 data pratinjau tersamar; endpoint live menunggu deployment commit server."],
        SHOT / "06-layar-kedua-obat-jadi-racikan.png", "Gambar 11. Layar kedua gabungan obat jadi/racikan dengan identitas tersamar")
    screenshot(doc, SHOT / "07-layar-tambahan-obat-jadi.png", "Gambar 12. Layar tambahan khusus obat jadi")
    screenshot(doc, SHOT / "08-layar-tambahan-racikan.png", "Gambar 13. Layar tambahan khusus racikan")

    section_text(doc, "11. Permintaan Pembelian (PR)",
        "PR mencatat kebutuhan sebelum memesan. Dasarnya dapat berupa stok minimum, tren konsumsi, kebutuhan program, penggantian recall, atau buffer yang disetujui. PR bukan bukti vendor boleh mengirim; ia harus melalui persetujuan dan dikonversi ke PO.",
        ["Buka Pengadaan → Permintaan Pembelian (PR) lalu Buat PR.", "Isi unit peminta, tanggal kebutuhan, alasan, dan prioritas.",
         "Tambahkan obat/barang dengan kuantitas, satuan, estimasi harga, dan spesifikasi.", "Lampirkan analisis kebutuhan atau stok bila diwajibkan.",
         "Simpan draft, review total, lalu ajukan persetujuan.", "Approver memeriksa kebutuhan, anggaran, duplikasi, dan kewajaran; setujui/tolak dengan catatan."],
        ["PR yang ditolak harus diperbaiki, bukan disalin tanpa alasan.", "UAT API menciptakan 50 PR; screenshot menu membuktikan formulir dan dashboard dapat dibuka.",
         "Angka dashboard dapat memakai horizon/filter berbeda dari batch UAT; audit volume mengandalkan respons API per kode UAT."],
        PROC / "02-pr-daftar-50.png", "Gambar 14. Dashboard PR dan akses ke daftar/formulir")
    screenshot(doc, PROC / "03-pr-formulir.png", "Gambar 15. Formulir PR")

    section_text(doc, "12. Pemesanan Pembelian (PO)",
        "PO merupakan komitmen kepada vendor setelah PR disetujui. Bedakan PO non-termin dan termin. Jadwal termin harus menjelaskan persentase/nilai, tanggal, syarat dokumen, serta hubungan dengan penerimaan agar pembayaran tidak lebih cepat atau lebih besar dari kewajiban.",
        ["Pilih Dari PR untuk mempertahankan jejak sumber; hindari membuat PO bebas bila kebijakan mewajibkan PR.",
         "Pilih vendor, alamat, mata uang, tanggal, target kirim, pajak, dan syarat.", "Periksa item, kuantitas, harga, diskon, PPN, dan total.",
         "Untuk termin, susun jadwal yang totalnya tepat 100% atau nilai penuh; untuk non-termin, tetapkan jatuh tempo tunggal.",
         "Ajukan persetujuan dan kirim PO hanya setelah status disetujui.", "Pantau sisa pesanan, keterlambatan, dan perubahan; amendment harus terdokumentasi."],
        ["Nomor PR dan PO harus dapat ditelusuri dua arah.", "UAT membuat 50 PO dengan pola termin/non-termin bergantian.",
         "Vendor, harga, dan termin memerlukan segregasi persetujuan."],
        PROC / "04-po-daftar-termin-nontermin.png", "Gambar 16. Dashboard Pemesanan Pembelian")
    screenshot(doc, PROC / "05-po-formulir-nontermin.png", "Gambar 17. Form PO non-termin")
    screenshot(doc, PROC / "06-po-formulir-termin.png", "Gambar 18. Form PO termin")

    section_text(doc, "13. BAST penerimaan barang/jasa",
        "BAST mengonfirmasi bahwa barang/jasa diterima dan menjadi dasar tagihan. Untuk obat, pemeriksaan mencakup nama, kekuatan, bentuk, jumlah, batch, ED, kondisi kemasan, izin/sertifikat bila relevan, suhu, dan kesesuaian PO.",
        ["Buka Penerimaan Barang/Jasa (BAST) dan pilih Dari PO.", "Pilih PO yang benar dan masukkan tanggal serta petugas penerima.",
         "Catat kuantitas diterima per baris, batch/ED, lokasi, dan hasil inspeksi.", "Jika parsial, sisakan kuantitas terbuka; jika berlebih atau tidak sesuai, jangan otomatis terima.",
         "Lampirkan surat jalan/foto/bukti suhu bila diwajibkan.", "Ajukan persetujuan; setelah sah, verifikasi status masuk stok dan nilai penerimaan."],
        ["BAST tidak boleh dibuat hanya dari faktur tanpa pemeriksaan fisik.", "UAT API membuat dan menyetujui 50 BAST.",
         "Sinkronisasi BAST ke modul kulakan lama pada demo masih memiliki constraint; lihat Bab 20."],
        PROC / "07-bast-daftar-50.png", "Gambar 19. Dashboard BAST")
    screenshot(doc, PROC / "08-bast-pilih-po.png", "Gambar 20. Pemilihan PO saat membuat BAST")

    section_text(doc, "14. Terima Tagihan Vendor",
        "Terima Tagihan mencatat faktur vendor setelah barang diterima. Terapkan three-way match antara PO, BAST, dan faktur. Perbedaan harga, kuantitas, pajak, termin, nomor faktur ganda, atau rekening vendor harus menahan proses hingga diselesaikan.",
        ["Buka Terima Tagihan Vendor dan pilih BAST terkait.", "Masukkan nomor/tanggal faktur, jatuh tempo, nilai dasar, pajak, potongan, dan total.",
         "Bandingkan baris faktur dengan PO dan kuantitas diterima di BAST.", "Periksa duplikasi nomor faktur untuk vendor yang sama.",
         "Lampirkan faktur dan bukti pajak; ajukan persetujuan.", "Setelah disetujui, pastikan kewajiban muncul pada daftar pembayaran vendor."],
        ["Jangan mengubah BAST agar sekadar sama dengan faktur.", "UAT API menghasilkan 50 tagihan vendor.",
         "Koreksi harus mempertahankan histori dan alasan."],
        PROC / "09-terima-tagihan-50.png", "Gambar 21. Dashboard Terima Tagihan Vendor")

    section_text(doc, "15. Pembayaran vendor dan Proses Transfer",
        "Pembayaran adalah mata rantai terakhir procure-to-pay. Petugas memilih tagihan disetujui, cara bayar, rekening sumber, jadwal, dan bukti. Setelah dibuat, proses transfer diperiksa dan disetujui oleh pihak berbeda bila kebijakan maker-checker berlaku.",
        ["Buka Proses Transfer → Pembayaran Vendor.", "Pilih tagihan yang telah disetujui dan belum lunas.",
         "Periksa vendor, rekening tujuan, rekening sumber, nilai, potongan, pajak, dan tanggal bayar.",
         "Buat pembayaran dan kelompokkan detail ke Proses Transfer.", "Reviewer membandingkan daftar transfer dengan dokumen sumber dan batas otorisasi.",
         "Setujui/eksekusi, simpan referensi bank, lalu tandai bukti dan status pelunasan.", "Serahkan daftar ke akuntansi untuk jurnal dan rekonsiliasi bank."],
        ["UAT: 50 pembayaran disetujui dan 50 detail tergabung pada Proses Transfer ID 10 senilai Rp45.730.000.",
         "Cara bayar Tunai dipetakan ke 111.101 KAS YAYASAN khusus skenario demo.", "Screenshot halaman dapat memakai filter yang tidak menampilkan kode batch; bukti volume berasal dari query/API terkontrol."],
        PROC / "10-pembayaran-vendor-50.png", "Gambar 22. Halaman Pembayaran Vendor pada Proses Transfer")
    screenshot(doc, diag["p2p"], "Gambar 23. Flow procure-to-pay obat")

    section_text(doc, "16. Pemetaan akun dan Jurnal Umum",
        "Pemetaan memakai daftar akun pada cetak_data_260904124814.xlsx (317 akun data). Untuk pembayaran vendor UAT, jurnal dibuat Dr 310.500 HUTANG VENDOR dan Cr 111.101 KAS YAYASAN. Persediaan, uang muka, PPN, pendapatan toko, dan HPP toko dipakai sesuai substansi transaksi serta kebijakan akuntansi institusi.",
        ["Buka master/pemetaan akun dan cari kode persis; cocokkan nama serta status aktif.", "Petakan vendor ke akun utang yang disetujui.",
         "Petakan metode bayar/rekening sumber ke kas/bank yang tepat.", "Buka Draft Jurnal untuk melihat kandidat dari dokumen operasional.",
         "Jika preview otomatis tersedia, periksa akun, debit, kredit, tanggal, deskripsi, dimensi, dan bukti.",
         "Jika preview belum tersedia, buat Jurnal Umum manual dari dokumen pembayaran: satu bukti, satu referensi unik, debit=credit.",
         "Simpan draft, review oleh pihak berwenang, lalu posting."],
        ["Jangan menebak atau menormalisasi kode akun bertanda koma pada sumber; minta konfirmasi pemilik data.",
         "UAT membuat dan memposting 50 jurnal manual pembayaran vendor secara idempoten.",
         "Jalur manual adalah kontrol sementara, bukan alasan membiarkan gap otomatis tanpa backlog."],
        PROC / "11-draft-jurnal-pengadaan.png", "Gambar 24. Draft Jurnal dan ringkasan aktivitas")
    screenshot(doc, diag["accounting"], "Gambar 25. Flow akuntansi pembayaran vendor")

    section_text(doc, "17. Posting dan laporan keuangan",
        "Posting mengubah draft menjadi catatan akuntansi yang menjadi dasar laporan. Sebelum posting, tanggal harus berada pada periode terbuka, akun aktif, debit sama dengan kredit, dokumen sumber valid, dan tidak ada duplikasi referensi. Setelah posting, koreksi dilakukan melalui pembalikan/adjustment sesuai kebijakan, bukan menghapus jejak.",
        ["Filter draft menurut periode dan kategori.", "Buka rincian dan cocokkan dengan pembayaran/BAST/faktur.", "Pastikan debit dan kredit seimbang serta dimensi organisasi benar.",
         "Posting oleh pengguna berwenang; simpan nomor jurnal.", "Buka Keseluruhan Jurnal lalu cari referensi.",
         "Periksa Buku Besar akun kas dan utang vendor; saldo harus bergerak berlawanan sesuai pembayaran.",
         "Periksa Neraca Saldo; total debit dan kredit harus seimbang.", "Tampilkan Laba Rugi, Neraca, dan Arus Kas untuk periode yang memuat transaksi.",
         "Ekspor bila perlu dan dokumentasikan rekonsiliasi."],
        ["UAT UI berhasil membuka enam laporan keuangan.", "Data laporan pada tangkapan tertentu dapat kosong karena filter periode/klasifikasi akun; ini tidak boleh disebut validasi angka penuh.",
         "Posting otomatis kategori pembayaran vendor dan fixed asset belum menyediakan preview pada demo."],
        ACC / "27-laporan-jurnal-umum.png", "Gambar 26. Keseluruhan Jurnal")
    for pth, cap in [
        (ACC / "28-laporan-buku-besar.png", "Gambar 27. Buku Besar"),
        (ACC / "29-laporan-neraca-saldo.png", "Gambar 28. Neraca Saldo"),
        (ACC / "24-laporan-laba-rugi.png", "Gambar 29. Laba Rugi"),
        (ACC / "25-laporan-neraca.png", "Gambar 30. Neraca"),
        (ACC / "26-laporan-arus-kas.png", "Gambar 31. Arus Kas"),
    ]:
        screenshot(doc, pth, cap)

    section_text(doc, "18. Rekonsiliasi, tutup hari, dan audit",
        "Tutup hari memastikan aktivitas farmasi, kas, stok, pengadaan, utang, dan jurnal memiliki posisi yang dapat dijelaskan. Jangan menganggap aplikasi yang tidak menampilkan error berarti rekonsiliasi selesai; bandingkan sumber independen dan dokumentasikan selisih.",
        ["Hitung kas fisik dan cocokkan dengan penerimaan menurut metode bayar.", "Rekonsiliasi jumlah transaksi dengan laporan penjualan dan bukti pembayaran.",
         "Rekonsiliasi mutasi item/batch dengan penerimaan PBF, penjualan, retur, opname, dan karantina.", "Cocokkan PR–PO–BAST–tagihan–pembayaran yang masih terbuka.",
         "Cocokkan daftar transfer dengan mutasi bank/kas.", "Cocokkan pembayaran vendor dengan jurnal terposting dan Buku Besar 310.500/111.101.",
         "Catat insiden, transaksi tertunda/offline, selisih, pemilik tindakan, dan target selesai.", "Supervisor menandatangani checklist dan mengunci periode sesuai kebijakan."],
        ["Selisih nol harus dibuktikan, bukan diasumsikan.", "Semua koreksi memiliki referensi ke dokumen asal.", "Log dan bukti menghindari data pasien yang tidak diperlukan."],
    )

    section_text(doc, "19. Bantuan dan penanganan masalah",
        "Setiap halaman menyediakan tombol Bantuan atau ikon tanda tanya. Buka bantuan sebelum mencoba perbaikan berisiko. Informasi bantuan menjelaskan tujuan, prasyarat, langkah, kontrol, dan masalah umum. Saat eskalasi, sertakan kode referensi, waktu, pengguna/peran, toko, versi, langkah reproduksi, dan tangkapan layar yang sudah disamarkan.",
        ["Tekan Bantuan untuk narasi menu; gunakan Tanya Jawab untuk pertanyaan kontekstual bila tersedia.",
         "Untuk koneksi: periksa internet, host/context, HTTPS, sertifikat, dan status server tanpa menghapus data lokal.",
         "Untuk menu hilang: periksa role/RBAC dan toko, jangan masuk memakai akun bersama.",
         "Untuk stok/batch: hentikan transaksi bila data meragukan, hitung fisik, periksa mutasi, lalu eskalasi.",
         "Untuk sinkronisasi tertunda: jangan mengulang bayar tanpa memastikan status server dan kode transaksi.",
         "Untuk jurnal gagal: baca preview/pesan, periksa akun/pemetaan/periode; gunakan jalur manual hanya dengan otorisasi.",
         "Untuk layar pasien: beralih ke panggilan manual dan lindungi privasi bila data basi atau koneksi putus."],
        ["Jangan menghapus cache/basis data sebagai percobaan pertama.", "Jangan mengubah jam perangkat untuk memanipulasi periode.",
         "Jangan mengirim token, kata sandi, atau data pasien dalam tiket."],
    )

    doc.add_heading("20. Batasan versi, prasyarat deploy, dan checklist go-live", level=1)
    note(doc, "Status penerimaan", "Build klien dan alur transaksi utama dapat dirilis untuk UAT lanjutan, tetapi penerimaan produksi bersifat bersyarat sampai endpoint antrean server dideploy, signing produksi tersedia, dan gap jurnal otomatis diputuskan pemilik proses.", AMBER)
    table(doc, ["Area", "Status v1.34.22", "Tindakan sebelum produksi"], [
        ("Penjualan obat jadi", "PASS — 50 transaksi API", "Pilot transaksi dan rekonsiliasi per toko"),
        ("Penebusan racikan", "PASS — 50 transaksi API, 100 resep terbaca", "Validasi klinis oleh apoteker lokal"),
        ("Penerimaan PBF", "PASS — batch tambahan 150 unit", "Tentukan relasi dengan BAST agar tidak double stock"),
        ("Layar pasien", "PASS komponen; endpoint demo belum tersedia", "Deploy AIS commit antrean dan uji live privacy/status"),
        ("Katalog volume", "2 live; generator 1000 committed belum deployed", "Deploy server lalu provision idempoten"),
        ("PR–PO–BAST–tagihan–bayar", "PASS API 50 per tahap", "Rekonsiliasi filter/dashboard dan dokumen lokal"),
        ("Posting jurnal manual", "PASS — 50 posted", "Rekonsiliasi Buku Besar dan approval"),
        ("Posting otomatis", "BLOCKED — preview belum tersedia", "Implementasi/aktifkan mapping atau sahkan SOP manual"),
        ("Kulakan generik", "BLOCKED untuk admin tanpa toko", "Uji dengan kasir/pedagang bertoko; perbaiki constraint sync BAST"),
        ("Android signing", "Build dapat memakai debug key", "Sediakan keystore produksi dan uji upgrade signature"),
    ], [3.4, 6.0, 7.6])
    doc.add_heading("Checklist go-live", level=2)
    for item in [
        "[ ] Release asset dan SHA-256 cocok; malware scan selesai.", "[ ] APK ditandatangani keystore produksi dan upgrade dari versi sebelumnya berhasil.",
        "[ ] Backup dan rollback diuji.", "[ ] Endpoint antrean live merespons dan mengirim IDENTITAS_DISAMARKAN.",
        "[ ] Katalog serta batch tiap toko lengkap; data-sample Tidak Aktif.", "[ ] Role kasir, apoteker, gudang, procurement, keuangan, akuntansi diuji least privilege.",
        "[ ] Printer, barcode, display kedua, dan failover jaringan diuji.", "[ ] COA dan pemetaan vendor/metode bayar ditandatangani pemilik akuntansi.",
        "[ ] Keputusan tertulis untuk jurnal otomatis vs SOP manual tersedia.", "[ ] Pilot end-to-end direkonsiliasi dan disetujui owner proses.",
    ]:
        bullet(doc, item)

    doc.add_heading("Lampiran A. Akun rujukan dari lampiran", level=1)
    doc.add_paragraph("Workbook CETAK DATA memuat 317 akun data. Pemetaan minimum yang digunakan/direkomendasikan dalam skenario UAT adalah:")
    table(doc, ["Kode", "Nama akun", "Penggunaan UAT"], [
        ("111.101", "KAS YAYASAN", "Kredit pembayaran vendor Tunai"),
        ("151.200", "PERSEDIAAN BARANG LAINNYA", "Persediaan obat bila kebijakan belum memakai subakun khusus"),
        ("171.200", "UANG MUKA PEMBELIAN", "Pembayaran di muka/termin sebelum hak tagih final"),
        ("310.500", "HUTANG VENDOR", "Debit saat membayar kewajiban vendor"),
        ("310.600", "UTANG USAHA TOKO", "Alternatif kewajiban toko bila ditetapkan owner"),
        ("310.301", "HUTANG PPN", "Pajak keluaran/utang PPN sesuai substansi"),
        ("410.900", "PENDAPATAN PENJUALAN TOKO", "Pendapatan penjualan apotek/toko sesuai kebijakan"),
        ("510.900", "BEBAN POKOK PENJUALAN TOKO", "HPP penjualan"),
    ], [3.0, 7.0, 7.0])
    note(doc, "Peringatan kualitas data", "Dua kode pada workbook memakai koma—112,102 CIMB dan 121,109 Bank Kaltim. Jangan menormalisasi menjadi titik tanpa persetujuan pemilik COA.", RED)

    doc.add_heading("Lampiran B. Kamus status dan checklist cepat", level=1)
    table(doc, ["Status", "Makna operasional", "Tindakan"], [
        ("MENUNGGU", "Belum mulai/menanti giliran", "Pastikan pemilik proses dan estimasi"),
        ("DISIAPKAN", "Obat sedang dikerjakan", "Jangan panggil pasien; update progres"),
        ("SIAP", "Pemeriksaan akhir selesai", "Panggil, verifikasi privat, edukasi, serahkan"),
        ("DRAFT", "Dokumen dapat diubah", "Review sebelum ajukan"),
        ("DIAJUKAN", "Menunggu otorisasi", "Approver memeriksa bukti"),
        ("DISETUJUI", "Sah untuk tahap berikut", "Proses sesuai segregasi tugas"),
        ("TERPOSTING", "Sudah menjadi jurnal", "Koreksi melalui reversal/adjustment"),
        ("BLOCKED", "Prasyarat belum terpenuhi", "Jangan dipaksa; catat owner dan kriteria retest"),
    ], [3.0, 7.0, 7.0])
    doc.add_heading("Checklist serah obat", level=2)
    for item in ["Benar pasien", "Benar obat/bahan", "Benar dosis/kekuatan", "Benar rute", "Benar waktu/frekuensi", "Benar jumlah", "Benar etiket", "Benar batch/ED", "Edukasi dipahami", "Dokumentasi selesai"]:
        bullet(doc, f"[ ] {item}")

    path = OUT / "Manual-Pengguna-POS-Apotik-v1.34.22.docx"
    doc.save(path)
    return path


def uat_report(diag):
    doc = Document()
    configure(doc, "Laporan UAT POS Apotik", "Laporan UAT")
    cover(doc, "User Acceptance Testing", "Laporan UAT POS Apotik",
          "Uji volume end-to-end: pelayanan farmasi, procure-to-pay, jurnal, dan laporan",
          "LULUS BERSYARAT — detail pengecualian wajib dibaca", BLUE)

    doc.add_heading("1. Ringkasan eksekutif", level=1)
    doc.add_paragraph(
        "UAT terhadap build Apotik 1.34.22+184 dan server demo.ecampus.id/ecampus membuktikan transaksi utama "
        "dapat dijalankan pada volume minimum 50 data per tahap. Sebanyak 50 penjualan obat jadi dan 50 "
        "penebusan racikan berhasil; 100 resep/racikan terbaca. Procure-to-pay menghasilkan 50 PR, 50 PO "
        "termin/non-termin, 50 BAST, 50 tagihan, dan 50 pembayaran disetujui. Satu proses transfer memuat "
        "50 detail senilai Rp45.730.000. Karena preview jurnal otomatis pembayaran vendor belum tersedia, "
        "50 jurnal umum manual dibuat secara idempoten dan semuanya terposting. Enam layar laporan akuntansi "
        "berhasil dibuka."
    )
    note(doc, "Keputusan", "LULUS BERSYARAT. Klien layak dipublikasikan untuk UAT/pilot, tetapi go-live penuh menunggu deploy endpoint antrean farmasi, keputusan gap posting otomatis, uji kasir bertoko untuk kulakan generik, dan keystore Android produksi.", AMBER)
    table(doc, ["Indikator", "Hasil"], [
        ("Penjualan obat jadi", "50/50 PASS"), ("Penebusan racikan", "50/50 PASS"),
        ("Resep/racikan terbaca", "100"), ("Penerimaan PBF", "150 unit, 1 batch baru PASS"),
        ("PR / PO / BAST / Tagihan / Pembayaran", "50 / 50 / 50 / 50 / 50 PASS API"),
        ("Proses transfer", "50 detail; Rp45.730.000; disetujui"),
        ("Jurnal pembayaran vendor", "50 manual, 50 terposting PASS"),
        ("Laporan", "6 menu terbuka PASS UI; angka perlu rekonsiliasi owner"),
        ("Layar antrean publik", "PASS komponen 50 data tersamar; BLOCKED integrasi endpoint demo"),
        ("Katalog", "2 live; target generator 1000 menunggu deploy server"),
    ], [7.5, 9.5])

    doc.add_heading("2. Tujuan, ruang lingkup, dan metode", level=1)
    doc.add_paragraph(
        "Tujuan UAT adalah memastikan alur yang diminta dapat dijalankan, data volume terbentuk, UI utama "
        "terbaca pada resolusi desktop, kontrol batch/privasi tersedia, dan hasil operasional dapat dilanjutkan "
        "ke akuntansi. Pengujian menggabungkan panggilan API terautentikasi untuk volume dan determinisme, "
        "integration_test Flutter Windows untuk widget/laporan/screenshot, serta inspeksi bukti. Pengujian ini "
        "tidak menggantikan validasi klinis, audit pajak, rekonsiliasi saldo awal, penetration test, atau "
        "sertifikasi perangkat rumah sakit."
    )
    table(doc, ["Komponen", "Lingkungan", "Bukti"], [
        ("Flutter Apotik", "Windows integration build, versi 1.34.22+184", "12 screenshot Apotik + log test pass"),
        ("AIS API", "https://demo.ecampus.id/ecampus/Api_eBisnis", "Respons API dan data berkode UAT"),
        ("Pengadaan", "Server demo + Flutter UI", "13 screenshot + 50 record per tahap"),
        ("Akuntansi", "Server demo + Flutter UI", "8 screenshot + 50 jurnal posted"),
        ("COA", "cetak_data_260904124814.xlsx", "317 akun; mapping kunci dan 2 anomali kode"),
    ], [4.0, 6.5, 6.5])
    screenshot(doc, diag["usecase"], "Gambar 1. Cakupan use case yang diuji")

    doc.add_heading("3. Prasyarat dan data uji", level=1)
    for item in [
        "Akun demo terautentikasi dengan hak Apotik, Pengadaan, Keuangan, dan Akuntansi yang diperlukan.",
        "Mode data sample diaktifkan hanya selama provisioning, kemudian dikembalikan ke Tidak Aktif setelah seluruh mutasi selesai.",
        "Toko uji ID 1; item UJI-PCT dan UJI-CDN tersedia. Katalog live memuat 2 item karena generator volume server baru belum terdeploy.",
        "Penerimaan PBF UAT-TERIMA-APT-13422-001 menambah 150 UJI-PCT ber-ED 2028-12-31.",
        "Prefix pengadaan UAT-VOL-PROC-APT13422-R4-20260904; proses transfer ID 10.",
        "Pemetaan pembayaran: Tunai → 111.101 KAS YAYASAN; vendor → 310.500 HUTANG VENDOR.",
    ]:
        bullet(doc, item)
    note(doc, "Integritas", "Kode transaksi penjualan deterministik UAT-APT-13422-JADI-001..050 dan UAT-APT-13422-RACIK-001..050. Run screenshot memakai mode capture-only yang membaca bukti ringkasan run API sebelumnya agar tidak menggandakan transaksi.")

    doc.add_heading("4. Matriks hasil UAT", level=1)
    cases = [
        ("APT-01", "Login, konfigurasi, role", "Identitas/peran/versi benar", "PASS", "Shell & konfigurasi API"),
        ("APT-02", "Dashboard", "50 resep menunggu terbaca", "PASS", "00-dashboard-operasional.png"),
        ("APT-03", "Katalog obat", ">=50 target", "BLOCKED ENV", "2 live; generator 1000 committed belum deployed"),
        ("APT-04", "Penerimaan PBF", "Tambah stok dan batch", "PASS", "150 unit, 1 batch; 05-penerimaan-pbf.png"),
        ("APT-05", "Obat jadi", "50 transaksi berhasil", "PASS", "50/50 API; 01-kasir-obat-jadi.png"),
        ("APT-06", "Resep/racikan", "50 transaksi + resep terbaca", "PASS", "50/50; 100 resep; 02-resep-racikan.png"),
        ("APT-07", "FEFO/expired", "Expired tak dipilih", "PASS", "Batch monitor dan pemilihan batch layak"),
        ("APT-08", "Laporan penjualan", "Status sukses, total tampil", "PASS", "Rp331.000; qty 107 pada screenshot"),
        ("APT-09", "Register terkendali", "Menu terbuka", "PASS UI", "10-register-obat-terkendali.png"),
        ("APT-10", "Laporan kedaluwarsa", "Status sukses", "PASS", "3 baris API; 11-laporan-kedaluwarsa.png"),
        ("APT-11", "Layar gabungan", "50 data tersamar tampil", "PASS KOMPONEN", "06-layar-kedua...png"),
        ("APT-12", "Endpoint antrean live", "list/simpan/status tersedia", "BLOCKED DEPLOY", "Server menjawab aksi belum tersedia"),
        ("P2P-01", "PR", "50 dibuat", "PASS API", "Prefix R4; halaman/form terbuka"),
        ("P2P-02", "PO termin/non-termin", "50 dibuat", "PASS API", "Pola bergantian"),
        ("P2P-03", "BAST", "50 dibuat/disetujui", "PASS API", "IDs proses 2..51"),
        ("P2P-04", "Terima Tagihan", "50 dibuat", "PASS API", "Prefix R4"),
        ("P2P-05", "Pembayaran", "50 disetujui", "PASS API", "Cara bayar Tunai; tanpa peringatan"),
        ("P2P-06", "Proses Transfer", "50 detail disetujui", "PASS", "ID 10; Rp45.730.000"),
        ("ACC-01", "Preview otomatis pembayaran", "Jurnal preview tersedia", "FAIL/BLOCKED", "pesanJurnal: Preview ... belum tersedia"),
        ("ACC-02", "Jurnal manual", "50 dibuat seimbang", "PASS", "Dr 310.500 / Cr 111.101"),
        ("ACC-03", "Posting jurnal", "50 terposting", "PASS", "Idempotent verify 50"),
        ("ACC-04", "6 laporan", "Layar dapat dibuka", "PASS UI", "Laba Rugi, Neraca, Arus Kas, Jurnal, Buku Besar, Neraca Saldo"),
        ("LEG-01", "Kulakan generik", "Admin demo dapat simpan", "BLOCKED", "Toko tidak diketahui; perlu akun pedagang bertoko"),
        ("SEC-01", "Data sample", "Nonaktif sesudah UAT", "PASS", "Diverifikasi portal: Tidak Aktif"),
    ]
    table(doc, ["ID", "Skenario", "Ekspektasi", "Status", "Bukti/catatan"], cases, [2.0, 4.0, 4.5, 2.7, 4.0])

    doc.add_heading("5. UAT pelayanan farmasi", level=1)
    doc.add_paragraph(
        "Run API menghasilkan tepat 50 transaksi obat jadi dan 50 racikan. Sebelum run, batch layak UJI-PCT "
        "tidak cukup untuk 100 unit sehingga penerimaan PBF 150 unit dilakukan. Setelah transaksi, layar kasir "
        "menunjukkan stok UJI-PCT 95 pada waktu screenshot dan laporan penjualan periode 30 hari menunjukkan "
        "UJI-PCT qty 105 senilai Rp315.000 serta UJI-CDN qty 2 senilai Rp16.000, total Rp331.000. Angka laporan "
        "mencakup data demo lain pada periodenya, sehingga bukan semata jumlah batch UAT."
    )
    screenshot(doc, SHOT / "00-dashboard-operasional.png", "Gambar 2. Dashboard: 50 resep menunggu")
    screenshot(doc, SHOT / "01-kasir-obat-jadi.png", "Gambar 3. Kasir obat jadi dan kontrol stok/LASA/terkendali")
    screenshot(doc, SHOT / "02-resep-racikan.png", "Gambar 4. Antrean resep")
    screenshot(doc, SHOT / "09-laporan-penjualan.png", "Gambar 5. Laporan penjualan setelah run")
    screenshot(doc, diag["finished"], "Gambar 6. Alur uji obat jadi")
    screenshot(doc, diag["compound"], "Gambar 7. Alur uji racikan")

    doc.add_heading("6. UAT layar pasien dan privasi", level=1)
    doc.add_paragraph(
        "Komponen produksi LayarAntreanFarmasiScreen diuji pada 1600×900 dengan 50 record terkontrol. Tiga "
        "mode—gabungan, khusus obat jadi, dan khusus racikan—berhasil dirender. Nama dan nomor rekam medis "
        "disamarkan, status dan loket terlihat, serta panel edukasi tampil. Namun endpoint antrean pada server "
        "demo memberi respons 'Aksi POS Apotik belum tersedia', sehingga hasil ini diklasifikasikan PASS KOMPONEN, "
        "bukan PASS integrasi. Retest wajib setelah commit server antrean terdeploy."
    )
    screenshot(doc, SHOT / "06-layar-kedua-obat-jadi-racikan.png", "Gambar 8. Bukti komponen layar gabungan dengan identitas tersamar")
    screenshot(doc, SHOT / "07-layar-tambahan-obat-jadi.png", "Gambar 9. Mode obat jadi")
    screenshot(doc, SHOT / "08-layar-tambahan-racikan.png", "Gambar 10. Mode racikan")

    doc.add_heading("7. UAT procure-to-pay", level=1)
    doc.add_paragraph(
        "Seed volume menyelesaikan 50 rangkaian PR–PO–BAST–tagihan–pembayaran. PO bergantian termin dan "
        "non-termin. Lima puluh pembayaran berstatus disetujui; detail DPC 36..85 dimasukkan dalam Proses "
        "Transfer ID 10 dan disetujui, total Rp45.730.000. Beberapa dashboard UI menampilkan agregat/filter "
        "yang berbeda dan tidak selalu memperlihatkan 50 kode UAT; karena itu bukti kuantitas berasal dari "
        "audit API, sementara screenshot membuktikan akses, formulir, dan rendering menu."
    )
    screenshot(doc, diag["p2p"], "Gambar 11. Flow procure-to-pay yang diuji")
    for pth, cap in [
        (PROC / "02-pr-daftar-50.png", "Gambar 12. Permintaan Pembelian"),
        (PROC / "03-pr-formulir.png", "Gambar 13. Form PR"),
        (PROC / "04-po-daftar-termin-nontermin.png", "Gambar 14. Pemesanan Pembelian"),
        (PROC / "06-po-formulir-termin.png", "Gambar 15. Form termin"),
        (PROC / "07-bast-daftar-50.png", "Gambar 16. Penerimaan BAST"),
        (PROC / "09-terima-tagihan-50.png", "Gambar 17. Terima Tagihan Vendor"),
        (PROC / "10-pembayaran-vendor-50.png", "Gambar 18. Pembayaran Vendor"),
    ]:
        screenshot(doc, pth, cap)

    doc.add_heading("8. UAT akuntansi dan laporan", level=1)
    doc.add_paragraph(
        "Kategori pembayaran termin vendor menghasilkan 50 draft pada ringkasan, tetapi rincian tidak "
        "menyediakan preview jurnal. Upaya posting kategori untuk 50 pembayaran dan 54 fixed asset tidak "
        "berhasil karena jurnal kosong dengan pesan bahwa preview belum tersedia. Untuk membuktikan alur "
        "akuntansi yang sah tanpa menyembunyikan gap, dibuat 50 Jurnal Umum manual—masing-masing mereferensikan "
        "pembayaran, debit 310.500 HUTANG VENDOR dan kredit 111.101 KAS YAYASAN—lalu seluruhnya terposting. "
        "Run kedua mengonfirmasi idempotensi: tidak membuat duplikat dan tetap menemukan 50 jurnal terposting."
    )
    note(doc, "Batas bukti laporan", "Enam halaman laporan berhasil dibuka. Beberapa screenshot menunjukkan belum ada data pada filter September/klasifikasi tertentu. Karena itu UAT membuktikan akses/rendering, sedangkan validasi angka laporan penuh tetap kriteria retest setelah klasifikasi COA dan periode disepakati.", AMBER)
    screenshot(doc, diag["accounting"], "Gambar 19. Jalur jurnal dengan fallback terkontrol")
    screenshot(doc, PROC / "11-draft-jurnal-pengadaan.png", "Gambar 20. Ringkasan Draft Jurnal")
    for pth, cap in [
        (ACC / "27-laporan-jurnal-umum.png", "Gambar 21. Keseluruhan Jurnal"),
        (ACC / "28-laporan-buku-besar.png", "Gambar 22. Buku Besar"),
        (ACC / "29-laporan-neraca-saldo.png", "Gambar 23. Neraca Saldo"),
        (ACC / "24-laporan-laba-rugi.png", "Gambar 24. Laba Rugi"),
        (ACC / "25-laporan-neraca.png", "Gambar 25. Neraca"),
        (ACC / "26-laporan-arus-kas.png", "Gambar 26. Arus Kas"),
    ]:
        screenshot(doc, pth, cap)

    doc.add_heading("9. Pemetaan akun dari workbook", level=1)
    doc.add_paragraph(
        "Workbook CETAK DATA memiliki 318 baris termasuk header, sehingga terdapat 317 akun data dan tidak "
        "ditemukan duplikasi kode pada pembacaan UAT. Pemetaan berikut dipakai sebagai rujukan; owner akuntansi "
        "tetap harus menyetujui klasifikasi akhir per institusi."
    )
    table(doc, ["Kode", "Nama", "Peran dalam UAT"], [
        ("111.101", "KAS YAYASAN", "Kredit pembayaran vendor"),
        ("151.200", "PERSEDIAAN BARANG LAINNYA", "Persediaan"),
        ("171.200", "UANG MUKA PEMBELIAN", "Uang muka"),
        ("310.500", "HUTANG VENDOR", "Debit pelunasan vendor"),
        ("310.600", "UTANG USAHA TOKO", "Alternatif utang toko"),
        ("310.301", "HUTANG PPN", "Kewajiban PPN"),
        ("410.900", "PENDAPATAN PENJUALAN TOKO", "Pendapatan"),
        ("510.900", "BEBAN POKOK PENJUALAN TOKO", "HPP"),
    ], [3.2, 8.0, 6.0])
    note(doc, "Anomali", "Baris kode 112,102 CIMB dan 121,109 Bank Kaltim memakai koma. Tidak ada normalisasi otomatis; pemilik COA harus memutuskan apakah itu penulisan sah atau salah input.", RED)

    doc.add_heading("10. Temuan, risiko, dan tindak lanjut", level=1)
    table(doc, ["ID", "Temuan", "Risiko", "Prioritas", "Owner/kriteria selesai"], [
        ("F-01", "Endpoint antrean belum terdeploy", "Layar tidak menerima status live", "P1", "DevOps AIS; list/simpan/status/hapus PASS di demo"),
        ("F-02", "Katalog live hanya 2", "Volume master tidak sesuai target", "P1", "Deploy generator 1000; provision; katalog >=1000"),
        ("F-03", "Preview otomatis pembayaran/fixed asset kosong", "Jurnal butuh kerja manual", "P1", "Akuntansi+backend; preview seimbang dan 50 auto-post"),
        ("F-04", "Kulakan generik admin tanpa toko", "Alur legacy tak selesai", "P2", "Uji akun pedagang bertoko; perbaiki constraint sync BAST"),
        ("F-05", "Filter/agregat UI tidak selalu menunjukkan 50 batch UAT", "Bukti visual membingungkan", "P2", "Tambah filter kode UAT/date dan counter hasil"),
        ("F-06", "APK mungkin debug-signed", "Upgrade/kepercayaan produksi", "P1", "Pemilik keystore; release-signed dan signature verified"),
        ("F-07", "2 kode akun bertanda koma", "Salah pemetaan bank", "P1", "Owner COA memberi keputusan tertulis"),
    ], [1.5, 4.7, 4.2, 1.5, 5.2])

    doc.add_heading("11. Kriteria retest dan penerimaan akhir", level=1)
    for item in [
        "Deploy server AIS yang memuat endpoint antrean dan generator demo 1000 obat + 1000 racikan.",
        "Provision ulang secara idempoten; buktikan katalog >=1000 dan data tidak mengganda pada run kedua.",
        "Buat 50 antrean live, ubah MENUNGGU→DISIAPKAN→SIAP, dan verifikasi tiga layar serta privasi IDENTITAS_DISAMARKAN.",
        "Selesaikan preview jurnal otomatis atau sahkan SOP manual; bila otomatis, 50 pembayaran harus menghasilkan jurnal seimbang tanpa duplikat.",
        "Uji kulakan generik dengan akun kasir/pedagang yang memiliki toko dan selesaikan sinkronisasi BAST.",
        "Validasi angka enam laporan pada periode yang benar terhadap 50 jurnal serta saldo awal yang disetujui.",
        "Build APK dengan keystore produksi dan uji fresh install, upgrade, login, transaksi, offline/sinkronisasi, serta rollback.",
        "Owner Farmasi, Procurement, Keuangan, Akuntansi, IT, dan Keamanan menandatangani penerimaan akhir.",
    ]:
        numbered(doc, item)
    table(doc, ["Peran penandatangan", "Nama", "Keputusan/tanggal", "Catatan"],
          [("Owner Farmasi", "", "", ""), ("Owner Procurement", "", "", ""),
           ("Owner Keuangan", "", "", ""), ("Owner Akuntansi", "", "", ""),
           ("IT/DevOps", "", "", ""), ("Keamanan/Privasi", "", "", "")], [4.5, 4.0, 4.0, 4.0])

    doc.add_heading("Lampiran. Inventaris bukti", level=1)
    evidence = []
    for group, directory in [("Apotik", SHOT), ("Pengadaan", PROC), ("Akuntansi", ACC), ("Diagram", DIAGRAM)]:
        for pth in sorted(directory.glob("*.png")):
            evidence.append((group, pth.name, f"{pth.stat().st_size:,} byte"))
    # Dua blok paralel menjaga inventaris lengkap tetap terbaca tanpa
    # menyisakan halaman terakhir yang hanya berisi satu-dua baris tabel.
    midpoint = (len(evidence) + 1) // 2
    paired_evidence = []
    for index in range(midpoint):
        left = evidence[index]
        right = evidence[index + midpoint] if index + midpoint < len(evidence) else ("", "", "")
        paired_evidence.append((*left, *right))
    inventory = table(
        doc,
        ["Kelompok", "Berkas", "Ukuran", "Kelompok", "Berkas", "Ukuran"],
        paired_evidence,
        [1.8, 5.5, 1.5, 1.8, 5.5, 1.5],
    )
    for row in inventory.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(6.5)

    path = OUT / "Laporan-UAT-POS-Apotik-v1.34.22.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    diag = diagrams()
    manual_path = manual(diag)
    uat_path = uat_report(diag)
    print(manual_path)
    print(uat_path)
